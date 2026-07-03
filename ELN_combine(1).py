# ==========================================
# 套件引入
# ==========================================
import win32com.client   # 連接 Windows Outlook COM 物件
import os
import re
import sys

APP_DIR = os.path.dirname(sys.executable) if getattr(sys, "frozen", False) else os.path.dirname(os.path.abspath(__file__))
BUNDLE_DIR = getattr(sys, "_MEIPASS", APP_DIR)
PROJECT_ROOT = r"D:\結構債流程優化專案(利用上手信) (T+1)"
ELN_FOLDER_ROOT = os.path.join(PROJECT_ROOT, "ELN folder")
REPORT_OUTPUT_DIR = os.path.join(ELN_FOLDER_ROOT, "ELN_Report")
FINAL_OUTPUT_DIR = os.path.join(ELN_FOLDER_ROOT, "ELN_Final_Output")
WORD_OUTPUT_DIR = os.path.join(ELN_FOLDER_ROOT, "Word_Output")


def resolve_runtime_path(filename):
    project_path = os.path.join(PROJECT_ROOT, filename)
    if os.path.exists(project_path):
        return project_path
    external_path = os.path.join(APP_DIR, filename)
    if os.path.exists(external_path):
        return external_path
    return os.path.join(BUNDLE_DIR, filename)


def ensure_directory(path):
    os.makedirs(path, exist_ok=True)
    return path


def build_report_output_path(start_date, end_date):
    return os.path.join(REPORT_OUTPUT_DIR, f"ELN_Report_{start_date}_to_{end_date}.xlsx")


def build_final_output_path(start_date, end_date):
    return os.path.join(FINAL_OUTPUT_DIR, f"ELN_Final_Output_{start_date}_to_{end_date}.xlsx")

import pandas as pd
import pdfplumber         # 解析 PDF 文字
import msoffcrypto        # 解密加密的 Office 檔案
import io
import datetime
from bs4 import BeautifulSoup   # 解析 HTML 內容
from dateutil import parser     # 智慧型日期字串解析
from docx import Document       # 產生 Word 文件
import html
import uuid
import shutil
import fitz                         # PyMuPDF，用於讀取 PDF 附件
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
import tempfile
import requests                              # Massive API 呼叫 (HSBC OBS Date 查詢)
import time                                  # API 速率限制等待

# ==========================================
# 全域設定常數
# ==========================================
OUTLOOK_ACCOUNT = "OPS1@feis.com.tw"        # 目標 Outlook 帳號
TARGET_FOLDER_NAME = "ELN到期與配息"          # 目標信箱子資料夾名稱
TEMP_PATH = os.path.join(tempfile.gettempdir(), "Temp_RPA_Process")           # 附件暫存路徑 (處理完後自動刪除)

# 各銀行加密 Excel 附件的解密密碼對應表 (key = 寄件者 Email domain 大寫)
PASSWORD_MAP = {
    "NOMURA": "1093834",
    "JPMORGAN": "feis_jpm",
    "CHASE": "feis_jpm",
}

# 已知上手寄件公司識別碼集合 (防呆用：偵測從未見過的新銀行/新寄件人)
KNOWN_COMPANIES = {
    "SGMARKETS", "BARCLAYS", "MORGAN", "MORGANSTANLEY",
    "HSBC", "NATIXIS", "BNP", "BNPPARIBAS",
    "NOMURA", "JPMORGAN", "CHASE", "BBVA", "UBS"
}

# 用於從信件主旨或內文判斷事件類型的關鍵字清單
EVENT_KEYWORDS = [
    "Coupon", "Interest Payment", "Redemption Payment", "Redemption",
    "Early Redemption Payment", "Early Termination", "Autocall",
    "Morgan Stanley Note Observation", "[MS] KO/Expiry Notification",
    "Events on your Structured Products", "ELN Fixings",
    "JPMorgan Securitized Trade", "Fixed Notification",
    "Fixing Events Notification",
    "Far Eastern International Securities",
    "Expiry", "Expired", "Expiration", "Expiration - Shares", "Maturity", "Final Redemption"
]

# ISIN 碼正規表達式：2 個大寫字母 + 9 個英數字 + 1 個數字，共 12 碼
ISIN_PATTERN = r"\b[A-Z]{2}[A-Z0-9]{9}\d\b"
# 日期正規表達式：支援 YYYY-MM-DD、DD/MM/YYYY 等格式
DATE_PATTERN_TEXT = r"\b\d{4}[/-]\d{2}[/-]\d{2}\b|\b\d{2}[/-]\d{2}[/-]\d{4}\b"

# 若暫存資料夾不存在則自動建立
if not os.path.exists(TEMP_PATH):
    os.makedirs(TEMP_PATH)

ensure_directory(REPORT_OUTPUT_DIR)
ensure_directory(FINAL_OUTPUT_DIR)
ensure_directory(WORD_OUTPUT_DIR)

# 步驟 4 草稿產生器專用暫存路徑
DRAFT_TEMP_PATH = os.path.join(tempfile.gettempdir(), "Temp_RPA_Process_DraftInline")
OUTPUT_DATA_PATH = os.path.join(DRAFT_TEMP_PATH, "output_data")
os.makedirs(DRAFT_TEMP_PATH, exist_ok=True)
os.makedirs(OUTPUT_DATA_PATH, exist_ok=True)

# 轉換股票履約價來源 Excel (第2欄=ISIN, 第4欄D=Ticker, 第7欄=Strike_Price)
ELN_EXTRACTED_DATA_PATH = resolve_runtime_path("ELN_Extracted_Data.xlsx")

# Massive API 設定 (用於 HSBC 換股贖回 OBS Date 查詢)
MASSIVE_BASE_URL = "https://api.massive.com/v2/aggs/ticker/{ticker}/range/1/day/{date_from}/{date_to}"
MASSIVE_DEFAULT_API_KEY = "EiQUqdBVV0jnnfZ4O53JZpTSxoooB0cS"
MASSIVE_FREE_TIER_INTERVAL_SECONDS = 12     # 免費方案每次呼叫間隔 (秒)
MASSIVE_OBS_SEARCH_WINDOW_DAYS = 60         # OBS Date 搜尋視窗 (往 Value Day 前推天數)


# ==========================================
# 共用工具函式
# ==========================================

def get_user_date_range():
    """互動式輸入日期區間，持續要求輸入直到格式正確且起始日 <= 結束日"""
    print("-" * 30)
    print("請輸入要查詢的信件日期區間 (格式: YYYY-MM-DD)")
    while True:
        try:
            start_str = input("起始日期 (例如 2023-10-01): ")
            end_str = input("結束日期 (例如 2023-10-31): ")
            s_date = datetime.datetime.strptime(start_str, "%Y-%m-%d").date()
            e_date = datetime.datetime.strptime(end_str, "%Y-%m-%d").date()
            if s_date > e_date:
                print("錯誤：起始日期不能晚於結束日期。")
                continue
            return s_date, e_date
        except ValueError:
            print("格式錯誤！請確保使用 YYYY-MM-DD 格式")


def get_outlook_messages(folder, start_date, end_date):
    """使用 Outlook Restrict 方法篩選指定日期區間的信件，若失敗則改為全量讀取"""
    messages = folder.Items
    messages.Sort("[ReceivedTime]", True)

    s_str = start_date.strftime("%m/%d/%Y 00:00 AM")
    e_str = (end_date + datetime.timedelta(days=1)).strftime("%m/%d/%Y 00:00 AM")

    filter_str = f"[ReceivedTime] >= '{s_str}' AND [ReceivedTime] < '{e_str}'"

    try:
        print(f"正在篩選郵件區間: {s_str} ~ {e_str} ...")
        return messages.Restrict(filter_str)
    except Exception as e:
        print(f"篩選失敗 (可能因版本問題)，改用全量讀取: {e}")
        return messages


def get_target_outlook_folder(outlook_namespace, account_name, target_folder_name):
    """在 Outlook 所有帳號中找到指定帳號，再依序定位收件匣及目標子資料夾"""
    target_store = None

    for i in range(1, outlook_namespace.Folders.Count + 1):
        root_folder = outlook_namespace.Folders.Item(i)
        try:
            folder_name = str(root_folder.Name).strip()
        except:
            continue

        print(f"偵測到 Outlook 帳號/信箱：{folder_name}")
        if folder_name.lower() == account_name.lower():
            target_store = root_folder
            break

    if target_store is None:
        raise Exception(f"找不到指定 Outlook 帳號：{account_name}")

    inbox = None
    for inbox_name in ["收件匣", "Inbox"]:
        try:
            inbox = target_store.Folders[inbox_name]
            print(f"成功找到收件匣：{inbox_name}")
            break
        except:
            pass

    if inbox is None:
        raise Exception(f"在帳號 {account_name} 底下找不到『收件匣』或『Inbox』")

    try:
        target_folder = inbox.Folders[target_folder_name]
    except Exception:
        available_folders = []
        try:
            for i in range(1, inbox.Folders.Count + 1):
                available_folders.append(inbox.Folders.Item(i).Name)
        except:
            pass
        raise Exception(
            f"在帳號 {account_name} 的收件匣底下找不到資料夾『{target_folder_name}』。"
            f"目前可見子資料夾：{available_folders}"
        )

    return target_folder


def extract_company_domain(email_address, sender_name):
    """從寄件者 Email domain 萃取公司識別碼 (例如 @sgmarkets.com → SGMARKETS)"""
    if not email_address:
        return sender_name.split()[0].upper()
    if "/O=" in email_address or "/o=" in email_address:
        return sender_name.split()[0].upper()

    match = re.search(r"@([a-zA-Z0-9-]+)\.", email_address)
    if match:
        domain = match.group(1).upper()
        if domain in ["GMAIL", "YAHOO", "HOTMAIL", "OUTLOOK"]:
            return sender_name.split()[0].upper()
        return domain
    else:
        return sender_name.split()[0].upper()


def get_event_type(text):
    """掃描文字中是否包含 EVENT_KEYWORDS 的關鍵字，回傳第一個符合者"""
    if not text:
        return "Unknown"
    for kw in EVENT_KEYWORDS:
        if kw.lower() in text.lower():
            return kw
    return "Unknown"


def convert_financial_date(date_str, is_us_format=False):
    """將各種日期字串統一轉換為 YYYY-MM-DD 格式"""
    if not date_str:
        return ""
    s = str(date_str).strip()
    try:
        use_dayfirst = False if is_us_format else True
        dt = parser.parse(s, fuzzy=True, dayfirst=use_dayfirst)
        return dt.strftime("%Y-%m-%d")
    except:
        return s


def normalize_html_column_name(col):
    """將 HTML 欄位名稱正規化：轉小寫、去換行/底線/連字號、合併多餘空白"""
    s = str(col).strip().lower()
    s = s.replace("\n", " ").replace("\r", " ")
    s = s.replace("_", " ").replace("-", " ")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def find_html_column(col_map, keywords):
    """在正規化欄位對應表中尋找包含任一關鍵字的欄位，回傳原始欄位名稱"""
    for norm_col, real_col in col_map.items():
        for kw in keywords:
            if kw in norm_col:
                return real_col
    return None


def clean_cell_text(val, default=""):
    """清理儲存格內容：NaN/None/空字串一律回傳 default"""
    if pd.isna(val):
        return default
    s = str(val).strip()
    if s.lower() in ["nan", "nat", "none", ""]:
        return default
    return s


def find_header_cell_index_by_keywords(header_cells, keyword_groups):
    for ci, cell in enumerate(header_cells):
        col_text = normalize_html_column_name(cell.get_text(" ", strip=True))
        for group in keyword_groups:
            if all(kw in col_text for kw in group):
                return ci
    return -1


def get_ticker_from_isin_yahoo(isin_code):
    """使用 Yahoo Finance Search API 透過 ISIN 碼查詢股票代號（symbol），失敗則回傳空字串"""
    try:
        url = f"https://query1.finance.yahoo.com/v1/finance/search?q={isin_code}&lang=en-US&region=US"
        resp = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        if resp.status_code == 200:
            data = resp.json()
            quotes = data.get("quotes", [])
            normalized_isin = str(isin_code or "").strip().upper()
            preferred_quote = ""
            for quote in quotes:
                symbol = str(quote.get("symbol") or "").strip()
                quote_isin = str(quote.get("isin") or "").strip().upper()
                quote_type = str(quote.get("quoteType") or "").strip().upper()
                if not symbol:
                    continue
                if quote_isin == normalized_isin and quote_type in {"EQUITY", "ETF"}:
                    return symbol
                if not preferred_quote and quote_type in {"EQUITY", "ETF"}:
                    preferred_quote = symbol
            if preferred_quote:
                return preferred_quote
            if quotes:
                return str(quotes[0].get("symbol") or "").strip()
    except Exception as e:
        print(f"[WARN] Yahoo Finance ISIN 查詢失敗 ({isin_code}): {e}")
    return ""


def extract_barclays_stock_ticker_from_text(stock_val):
    stock_text = clean_cell_text(stock_val, "")
    if not stock_text or stock_text == "-":
        return ""
    isin_match = re.search(ISIN_PATTERN, stock_text.upper())
    if not isin_match:
        return ""
    return get_ticker_from_isin_yahoo(isin_match.group(0))


def looks_like_isin_only(text):
    cleaned = str(text or "").strip().upper()
    if not cleaned:
        return False
    cleaned = re.sub(r"^換股標的:\s*", "", cleaned)
    cleaned = re.sub(r"\([^)]*\)", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        return False
    first_token = cleaned.split()[0]
    return bool(re.fullmatch(r"[A-Z]{2}[A-Z0-9]{9}\d", first_token))


def looks_like_isin_only(text):
    cleaned = str(text or "").strip().upper()
    if not cleaned:
        return False
    if ":" in cleaned:
        cleaned = cleaned.split(":", 1)[1].strip()
    cleaned = re.sub(r"\([^)]*\)", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        return False
    first_token = cleaned.split()[0]
    return bool(re.fullmatch(r"[A-Z]{2}[A-Z0-9]{9}\d", first_token))


def _lookup_obs_date_from_yahoo(underlying_ticker, value_day, closing_price):
    """
    使用 Yahoo Finance 歷史收盤價查詢 OBS Date (Natixis FINAL 換股轉換專用)。
    Bloomberg ticker (如 'NCLH UN') → 取第一段作為 Yahoo Finance symbol。
    在 Value Day 前 MASSIVE_OBS_SEARCH_WINDOW_DAYS 天範圍內比對收盤價。
    回傳 YYYY-MM-DD str，查無或異常則回傳 None。
    """
    symbol = str(underlying_ticker or "").strip().split()[0].upper()
    if not symbol or not value_day or closing_price is None:
        return None
    try:
        vd = datetime.datetime.strptime(value_day, "%Y-%m-%d").date()
    except Exception:
        return None
    date_from = vd - datetime.timedelta(days=30)
    _epoch = datetime.datetime(1970, 1, 1)
    start_ts = int((datetime.datetime(date_from.year, date_from.month, date_from.day) - _epoch).total_seconds())
    end_ts   = int((datetime.datetime(vd.year, vd.month, vd.day, 23, 59, 59) - _epoch).total_seconds())
    url = f"https://query1.finance.yahoo.com/v8/finance/chart/{symbol}"
    params  = {"period1": start_ts, "period2": end_ts, "interval": "1d", "events": "history"}
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        print(f"  [Yahoo OBS] 查詢 {symbol} [{date_from}~{vd}] Closing={closing_price}")
        resp = requests.get(url, params=params, headers=headers, timeout=15)
        resp.raise_for_status()
        data   = resp.json()
        result = data.get("chart", {}).get("result", [])
        if not result:
            print(f"  [Yahoo OBS] ✗ 無資料 ({symbol})")
            return None
        timestamps = result[0].get("timestamp", [])
        quotes     = result[0].get("indicators", {}).get("quote", [{}])
        closes     = quotes[0].get("close", []) if quotes else []
        target_rounded = round(float(closing_price), 2)
        candidates = []
        for ts, c in zip(timestamps, closes):
            if c is None:
                continue
            if round(float(c), 2) == target_rounded:
                try:
                    bar_date = (_epoch + datetime.timedelta(seconds=ts)).date()
                    if bar_date <= vd:
                        candidates.append(bar_date)
                except Exception:
                    continue
        if not candidates:
            print(f"  [Yahoo OBS] ✗ 未找到符合 Closing={closing_price} 的日期")
            return None
        obs_date = max(candidates).strftime("%Y-%m-%d")
        print(f"  [Yahoo OBS] ✓ OBS Date={obs_date}")
        return obs_date
    except Exception as e:
        print(f"  [Yahoo OBS] API 錯誤: {e}")
        return None


# ==========================================
# === 步驟 1：ELN1 — Outlook 信件擷取 ===
# ==========================================

def process_sg_markets_html(html_content):
    # SG Markets 信件包含兩張 HTML 表格：
    #   ① 換股明細表 (含 Underlying Ticker、每張票股數、零股現金)
    #   ② 主事件表  (含事件類型、Payment Date、Valuation Date)
    # 策略：先掃描所有表找換股明細存入 isin_details_map，再讀主事件表時對應合併
    extracted_data = []

    try:
        dfs = pd.read_html(io.StringIO(html_content))
        isin_details_map = {}  # key=ISIN，value={ticker, shares, cash}

        for df_idx, df in enumerate(dfs, start=1):
            try:
                df.columns = [str(c).strip() for c in df.columns]
                norm_map = {normalize_html_column_name(c): c for c in df.columns}

                target_isin_col = find_html_column(norm_map, ["isin"])
                target_ticker_col = find_html_column(norm_map, ["underlying ticker"])
                share_col = find_html_column(norm_map, ["number of shares per note"])
                cash_col = find_html_column(norm_map, ["residual amount per note"])

                if target_isin_col and target_ticker_col:
                    for _, row in df.iterrows():
                        isin_val = clean_cell_text(row.get(target_isin_col, ""))
                        if not isin_val:
                            continue

                        ticker_val = clean_cell_text(row.get(target_ticker_col, ""))
                        sh_val = clean_cell_text(row.get(share_col, "0"), default="0") if share_col else "0"
                        c_val = clean_cell_text(row.get(cash_col, "0"), default="0") if cash_col else "0"

                        isin_details_map[isin_val] = {
                            "ticker": ticker_val,
                            "shares": sh_val,
                            "cash": c_val
                        }
            except Exception:
                continue

        for df_idx, df in enumerate(dfs, start=1):
            try:
                df.columns = [str(c).strip() for c in df.columns]
                norm_map = {normalize_html_column_name(c): c for c in df.columns}

                if find_html_column(norm_map, ["underlying ticker"]) or find_html_column(norm_map, ["delivered underlying"]):
                    continue

                isin_col = find_html_column(norm_map, ["isin"])
                event_col = find_html_column(norm_map, ["event type"])
                pay_date_col = find_html_column(norm_map, ["payment date"])
                obs_date_col = find_html_column(norm_map, ["valuation date"])

                if not (isin_col and event_col):
                    continue

                for row_idx, row in df.iterrows():
                    isin = clean_cell_text(row.get(isin_col, ""))
                    if not isin or len(isin) <= 5:
                        continue

                    raw_event = row.get(event_col, "")
                    raw_pay_date = row.get(pay_date_col, "") if pay_date_col else ""
                    raw_obs = row.get(obs_date_col, "") if obs_date_col else ""

                    e_type = clean_cell_text(raw_event, default="Unknown")

                    pay_date = ""
                    if pay_date_col and pd.notna(raw_pay_date):
                        pay_date = convert_financial_date(str(raw_pay_date).strip())

                    obs_date = ""
                    if obs_date_col and pd.notna(raw_obs):
                        obs_date = convert_financial_date(str(raw_obs).strip())

                    sh_result = "0"
                    c_result = "0"
                    if isin in isin_details_map:
                        sh_result = isin_details_map[isin]["shares"]
                        c_result = isin_details_map[isin]["cash"]

                    display_value = f"Event:{e_type}, Date:{pay_date}"
                    if "MATURITY" in e_type.upper() and isin in isin_details_map:
                        underlying_ticker = isin_details_map[isin]["ticker"]
                        if underlying_ticker:
                            display_value = f"換股標的: {underlying_ticker}"

                    extracted_data.append({
                        "ISIN": isin,
                        "Event Type": e_type,
                        "Value Day": pay_date,
                        "Obs Day": obs_date,
                        "Value": display_value,
                        "Source": "SG Markets HTML",
                        "Denom": 10000,
                        "Strike": "0",
                        "Shares": sh_result,
                        "FractionalCash": c_result
                    })
            except Exception:
                continue

    except Exception as e:
        print(f"SG Markets HTML 解析錯誤: {e}")

    return extracted_data


def process_bnp_pdf(file_path, subject):
    """BNP Paribas：以 pdfplumber 擷取全頁文字後用 regex 搜尋關鍵欄位"""
    extracted_data = []
    full_text = ""

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"

        isin_val = "Unknown"
        found_isin = re.search(ISIN_PATTERN, full_text)
        if found_isin:
            isin_val = found_isin.group(0)

        subj_lower = subject.lower()
        event_type = "Unknown"
        if "expiry" in subj_lower:
            event_type = "Expiry"
        elif "early termination" in subj_lower:
            event_type = "Early Termination"
        elif "redemption" in subj_lower:
            event_type = "Redemption"
        elif "coupon" in subj_lower:
            event_type = "Coupon"
        else:
            event_type = get_event_type(subject)

        pay_date = ""
        obs_date = ""

        date_match = re.search(
            r"Payment Date[\s\S]{0,30}?(\d{1,2}[- ]?[A-Za-z]{3}[- ]?\d{2,4})",
            full_text,
            re.IGNORECASE
        )
        if date_match:
            pay_date = convert_financial_date(date_match.group(1))

        obs_date_match = re.search(
            r"(?<!Trade )(?<!Maturity )(?<!Redemption )Date\s*:\s*(\d{1,2}[- ]?[A-Za-z]{3,}[- ]?\d{2,4})",
            full_text,
            re.IGNORECASE
        )
        if obs_date_match:
            obs_date = convert_financial_date(obs_date_match.group(1))

        val_raw = "See PDF"
        if event_type == "Expiry" or "REDEMPTION" in event_type.upper():
            share_match = re.search(
                r"Share Name\s+(.*?)(?=\s*Share Reference|\s*Share amount|\n)",
                full_text,
                re.IGNORECASE
            )
            if share_match:
                ticker = share_match.group(1).strip()
                ticker = re.split(r"Share\s*Reference", ticker, flags=re.IGNORECASE)[0]
                ticker = ticker.strip()
                if len(ticker) < 30:
                    val_raw = f"換股標的: {ticker}"

        denom_val = "10000"
        shares_val = "0"
        cash_val = "0"
        strike_val = "0"

        if "REDEMPTION" in event_type.upper() or "EXPIRY" in event_type.upper():
            s_match = re.search(
                r"Share amount per Notional Amount of each Certificate\s*(\d+(?:\.\d+)?)",
                full_text,
                re.IGNORECASE
            )
            c_match = re.search(
                r"Residual cash for fractional entitlement\s*([\d,]+(?:\.\d+)?)\s*USD",
                full_text,
                re.IGNORECASE
            )
            n_match = re.search(
                r"Notional Amount:\s*([\d,]+(?:\.\d+)?)\s*USD",
                full_text,
                re.IGNORECASE
            )

            if n_match:
                denom_val = n_match.group(1).replace(",", "")
            if s_match:
                shares_val = s_match.group(1)
            if c_match:
                cash_val = c_match.group(1).replace(",", "")

        if isin_val != "Unknown" or pay_date:
            extracted_data.append({
                "ISIN": isin_val,
                "Event Type": event_type,
                "Value Day": pay_date,
                "Obs Day": obs_date,
                "Value": val_raw,
                "Source": "BNP PDF",
                "Denom": denom_val,
                "Strike": strike_val,
                "Shares": shares_val,
                "FractionalCash": cash_val
            })

    except Exception as e:
        print(f"BNP PDF 解析錯誤: {e}")

    return extracted_data


def process_hsbc_html(html_content, subject=""):
    """HSBC 一般配息/贖回信：解析 HTML 表格，定位 Payment Date / ISIN / Observation Date"""
    extracted_data = []
    soup = BeautifulSoup(html_content, 'html.parser')
    target_header = soup.find(string=re.compile(r"Payment\s*Date", re.IGNORECASE))

    if target_header:
        header_cell = target_header.find_parent(['th', 'td'])
        header_row = header_cell.find_parent('tr')
        table = header_row.find_parent('table')

        if header_row and table:
            header_cells = header_row.find_all(['th', 'td'])
            try:
                date_col_index = header_cells.index(header_cell)
            except:
                date_col_index = -1

            isin_col_index = -1
            isin_header = header_row.find(string=re.compile(r"ISIN", re.IGNORECASE))
            if isin_header:
                isin_cell = isin_header.find_parent(['th', 'td'])
                if isin_cell in header_cells:
                    isin_col_index = header_cells.index(isin_cell)

            obs_col_index = -1
            obs_header = header_row.find(string=re.compile(r"Observation\s*date", re.IGNORECASE))
            if obs_header:
                obs_cell = obs_header.find_parent(['th', 'td'])
                if obs_cell in header_cells:
                    obs_col_index = header_cells.index(obs_cell)

            subj_lower = subject.lower()
            if "early redemption" in subj_lower:
                event_val = "Early Redemption"
            elif "redemption" in subj_lower:
                event_val = "Redemption"
            else:
                event_val = "Interest Payment"

            all_rows = table.find_all('tr')
            try:
                start_row_idx = all_rows.index(header_row) + 1
            except:
                start_row_idx = 1

            for row in all_rows[start_row_idx:]:
                cells = row.find_all(['td', 'th'])
                if len(cells) > date_col_index:
                    raw_date = cells[date_col_index].get_text(strip=True)
                    clean_date = convert_financial_date(raw_date)

                    clean_obs_date = ""
                    if obs_col_index != -1 and len(cells) > obs_col_index:
                        raw_obs = cells[obs_col_index].get_text(strip=True)
                        clean_obs_date = convert_financial_date(raw_obs)

                    isin_val = "Unknown"
                    if isin_col_index != -1 and len(cells) > isin_col_index:
                        isin_val = cells[isin_col_index].get_text(strip=True)

                    if not re.search(ISIN_PATTERN, isin_val):
                        row_text = row.get_text(" ", strip=True)
                        found_isin = re.search(ISIN_PATTERN, row_text)
                        if found_isin:
                            isin_val = found_isin.group(0)

                    if clean_date and isin_val and len(isin_val) > 5:
                        extracted_data.append({
                            "ISIN": isin_val,
                            "Event Type": event_val,
                            "Value Day": clean_date,
                            "Obs Day": clean_obs_date,
                            "Value": f"HSBC Table: {event_val}",
                            "Source": "HSBC HTML Table",
                            "Denom": "10000",
                            "Strike": "0",
                            "Shares": "0",
                            "FractionalCash": "0"
                        })

    return extracted_data


def process_hsbc_redemption_shares(text):
    """HSBC 換股贖回通知 (純文字)：以 ISIN 切段後在各段搜尋換股資訊。"""
    results = []

    note_iter = re.finditer(r"ISIN\s+([A-Z]{2}[A-Z0-9]{9}\d)", text)
    matches = list(note_iter)

    for i, match in enumerate(matches):
        note_isin = match.group(1)
        start_pos = match.end()
        end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        segment = text[start_pos:end_pos]

        stock_match = re.search(r"\(([^)]+)\)\s*\(([A-Z]{2}[A-Z0-9]{9}\d)\)", segment)
        stock_ticker = "Unknown"
        stock_isin = "Unknown"
        if stock_match:
            stock_ticker = stock_match.group(1).strip()
            stock_isin = stock_match.group(2).strip()

        date_match = re.search(r"proceeds value\s+(\d{1,2}[A-Za-z]{3}\d{2,4})", segment, re.IGNORECASE)
        val_date = "Unknown"
        if date_match:
            val_date = convert_financial_date(date_match.group(1))

        denom_val = "10000"
        shares_val = "0"
        cash_val = "0"
        strike_val = "0"

        d_match = re.search(r"per denomination of USD\s*([\d,]+(?:\.\d+)?)", segment, re.IGNORECASE)
        s_match = re.search(r"Shares per Note is\s*(\d+(?:\.\d+)?)", segment, re.IGNORECASE)
        c_match = re.search(r"Fractional entitlement settle by cash in USD\s*([\d,]+(?:\.\d+)?)", segment, re.IGNORECASE)
        st_match = re.search(
            fr"strike.*?{re.escape(note_isin)}[\s\r\n]+(?:\d+(?:\.\d+)?)[\s\r\n]+([\d,]+(?:\.\d+)?)",
            text,
            re.IGNORECASE | re.DOTALL
        )

        if d_match:
            denom_val = d_match.group(1).replace(",", "")
        if s_match:
            shares_val = s_match.group(1)
        if c_match:
            cash_val = c_match.group(1).replace(",", "")
        if st_match:
            strike_val = st_match.group(1)

        results.append({
            "ISIN": note_isin,
            "Event Type": "Redemption in Shares",
            "Value Day": val_date,
            "Obs Day": "N/A",
            "Value": f"換股標的: {stock_ticker} ({stock_isin})",
            "Source": "HSBC Share Redem Body",
            "Denom": denom_val,
            "Strike": strike_val,
            "Shares": shares_val,
            "FractionalCash": cash_val,
        })

    return results


def process_generic_html_table(html_content, target_keywords, source_name, is_us_format=False):
    """通用 HTML 表格解析函式，適用於 Morgan Stanley、Barclays、Natixis"""
    extracted_data = []
    soup = BeautifulSoup(html_content, 'html.parser')
    target_header = None

    for kw in target_keywords:
        target_header = soup.find(string=re.compile(kw, re.IGNORECASE))
        if target_header:
            break

    if target_header:
        header_cell = target_header.find_parent(['th', 'td'])
        header_row = header_cell.find_parent('tr')
        table = header_row.find_parent('table')

        if header_row and table:
            header_cells = header_row.find_all(['th', 'td'])
            try:
                date_col_index = header_cells.index(header_cell)
            except:
                date_col_index = -1

            isin_col_index = -1
            isin_header = header_row.find(string=re.compile(r"ISIN", re.IGNORECASE))
            if isin_header:
                isin_cell = isin_header.find_parent(['th', 'td'])
                if isin_cell in header_cells:
                    isin_col_index = header_cells.index(isin_cell)

            event_col_index = -1
            event_header = header_row.find(string=re.compile(r"Event|Observation\s*Type", re.IGNORECASE))
            if event_header:
                event_cell = event_header.find_parent(['th', 'td'])
                if event_cell in header_cells:
                    event_col_index = header_cells.index(event_cell)

            obs_col_index = -1
            obs_header = header_row.find(string=re.compile(r"Fixing\s*Date|Valuation\s*Date", re.IGNORECASE))
            if obs_header:
                obs_cell = obs_header.find_parent(['th', 'td'])
                if obs_cell in header_cells:
                    obs_col_index = header_cells.index(obs_cell)

            ticker_col_index = find_header_cell_index_by_keywords(
                header_cells,
                [["shares", "to", "be", "delivered"], ["share", "element"]]
            )

            share_col_index = -1
            share_header = header_row.find(string=re.compile(r"Total\s*Share\s*Deliver.*|Quantity\s*of\s*stock/denom.*", re.IGNORECASE))
            if share_header:
                share_cell = share_header.find_parent(['th', 'td'])
                if share_cell in header_cells:
                    share_col_index = header_cells.index(share_cell)

            denom_col_index = -1
            denom_header = header_row.find(string=re.compile(r"Denomination", re.IGNORECASE))
            if denom_header:
                denom_cell = denom_header.find_parent(['th', 'td'])
                if denom_cell in header_cells:
                    denom_col_index = header_cells.index(denom_cell)

            ms_total_payable_index = -1
            mstp_header = header_row.find(string=re.compile(r"Total\s*Amount\s*Payable", re.IGNORECASE))
            if mstp_header and mstp_header.find_parent(['th', 'td']) in header_cells:
                ms_total_payable_index = header_cells.index(mstp_header.find_parent(['th', 'td']))

            ms_coupon_index = -1
            msc_header = header_row.find(string=re.compile(r"Coupon\s*Amount\s*per\s*Denom", re.IGNORECASE))
            if msc_header and msc_header.find_parent(['th', 'td']) in header_cells:
                ms_coupon_index = header_cells.index(msc_header.find_parent(['th', 'td']))

            cash_col_index = -1
            cash_header = header_row.find(string=re.compile(r"Redemption\s*Amount\s*per\s*Denom|Redemption\s*Amt/Denom", re.IGNORECASE))
            if cash_header:
                cash_cell = cash_header.find_parent(['th', 'td'])
                if cash_cell in header_cells:
                    cash_col_index = header_cells.index(cash_cell)

            barclays_stock_col_index = find_header_cell_index_by_keywords(
                header_cells,
                [["stock", "redemption", "sedol", "isin"]]
            )

            # Natixis FINAL 換股轉換專用欄位索引 (使用 get_text 應對多行 <th> 文字)
            natixis_share_delivery_col_index = -1
            natixis_final_price_col_index = -1
            natixis_nb_shares_col_index = -1
            natixis_residual_cash_col_index = -1
            if "natixis" in source_name.lower():
                for ci, hcell in enumerate(header_cells):
                    col_text = re.sub(r'\s+', ' ', hcell.get_text(' ', strip=True)).strip().upper()
                    if re.search(r'SHARE\s+DELIVERY', col_text):
                        natixis_share_delivery_col_index = ci
                    elif re.search(r'FINAL\s+PRICE', col_text):
                        natixis_final_price_col_index = ci
                    elif re.search(r'NB\s+SHARE[S]?\s+PER\s+NOTE', col_text):
                        natixis_nb_shares_col_index = ci
                    elif re.search(r'RESIDUAL\s+CASH\s+PER\s+NOTE', col_text):
                        natixis_residual_cash_col_index = ci

            all_rows = table.find_all('tr')
            try:
                start_row_idx = all_rows.index(header_row) + 1
            except:
                start_row_idx = 1

            for row in all_rows[start_row_idx:]:
                cells = row.find_all(['td', 'th'])
                if len(cells) > date_col_index:
                    raw_date = cells[date_col_index].get_text(strip=True)
                    clean_date = convert_financial_date(raw_date, is_us_format=is_us_format)

                    clean_obs_date = ""
                    if obs_col_index != -1 and len(cells) > obs_col_index:
                        raw_obs = cells[obs_col_index].get_text(strip=True)
                        clean_obs_date = convert_financial_date(raw_obs, is_us_format=is_us_format)

                    isin_val = "Unknown"
                    if isin_col_index != -1 and len(cells) > isin_col_index:
                        isin_val = cells[isin_col_index].get_text(strip=True)

                    if not re.search(ISIN_PATTERN, isin_val):
                        row_text = row.get_text(" ", strip=True)
                        found_isin = re.search(ISIN_PATTERN, row_text)
                        if found_isin:
                            isin_val = found_isin.group(0)

                    event_val = "Event"
                    if event_col_index != -1 and len(cells) > event_col_index:
                        event_val = cells[event_col_index].get_text(strip=True)

                    display_val = f"Table: {event_val}"

                    share_val = "0"
                    if share_col_index != -1 and len(cells) > share_col_index:
                        s_text = cells[share_col_index].get_text(strip=True).replace(",", "")
                        if s_text and s_text != "-":
                            share_val = s_text

                    ticker_val = ""
                    if ticker_col_index != -1 and len(cells) > ticker_col_index:
                        ticker_val = cells[ticker_col_index].get_text(strip=True)

                    # Morgan Stanley Expiration - Shares 特殊處理
                    if "expiration" in event_val.lower() and "shares" in event_val.lower():
                        if ticker_val and ticker_val != "-":
                            display_val = f"換股標的:{ticker_val}"
                        # 確保將 Expiration - Shares 視為換股事件
                        if not display_val.startswith("換股標的:"):
                            display_val = f"換股標的: {event_val}"

                    # Barclays Redemption 換股處理：從 Stock Redemption Sedol/ISIN 欄取得 ISIN，
                    # 再透過 Yahoo Finance Search API 查詢股票名稱
                    if "barclays" in source_name.lower() and "redemption" in event_val.lower():
                        if barclays_stock_col_index != -1 and len(cells) > barclays_stock_col_index:
                            stock_val = cells[barclays_stock_col_index].get_text(strip=True)
                            if stock_val and stock_val != "-" and len(stock_val) > 1:
                                yahoo_ticker = extract_barclays_stock_ticker_from_text(stock_val)
                                if yahoo_ticker:
                                    stock_val = yahoo_ticker
                                # 若欄位內容符合 ISIN 格式，查詢 Yahoo Finance 取得股票名稱
                                isin_match = re.search(ISIN_PATTERN, stock_val)
                                if isin_match:
                                    found_isin = isin_match.group(0)
                                    country_code = found_isin[:2]          # 例：US、HK、TW
                                    yahoo_ticker = get_ticker_from_isin_yahoo(found_isin)
                                    if yahoo_ticker:
                                        display_val = f"換股標的:{yahoo_ticker} {country_code}"
                                    else:
                                        display_val = f"換股標的:{stock_val}"
                                else:
                                    display_val = f"換股標的:{stock_val}"

                    denom_val = "10000"
                    if denom_col_index != -1 and len(cells) > denom_col_index:
                        d_text = cells[denom_col_index].get_text(strip=True).replace(",", "")
                        if d_text:
                            denom_val = d_text

                    cash_val = "0"
                    c_texts = []

                    if ms_total_payable_index != -1 and len(cells) > ms_total_payable_index:
                        v = cells[ms_total_payable_index].get_text(strip=True).replace(",", "")
                        if v and v != "-":
                            c_texts.append(v)

                    if ms_coupon_index != -1 and len(cells) > ms_coupon_index:
                        v = cells[ms_coupon_index].get_text(strip=True).replace(",", "")
                        if v and v != "-":
                            c_texts.append(f"(+Coupon:{v})")

                    if len(c_texts) > 0:
                        cash_val = " ".join(c_texts)
                    elif cash_col_index != -1 and len(cells) > cash_col_index:
                        c_text = cells[cash_col_index].get_text(strip=True).replace(",", "")
                        if c_text and c_text != "-":
                            cash_val = c_text

                    # Natixis FINAL 換股轉換：覆寫換股標的 / 接零股 / 接現金，並提取收盤價供 OBS Date 查詢
                    natixis_closing_price = None
                    if "natixis" in source_name.lower():
                        if natixis_share_delivery_col_index != -1 and len(cells) > natixis_share_delivery_col_index:
                            nd_ticker = cells[natixis_share_delivery_col_index].get_text(strip=True)
                            if nd_ticker and nd_ticker not in ("-", ""):
                                ticker_val = nd_ticker
                                display_val = f"換股標的:{nd_ticker}"
                        if natixis_nb_shares_col_index != -1 and len(cells) > natixis_nb_shares_col_index:
                            nbs_text = cells[natixis_nb_shares_col_index].get_text(strip=True).replace(",", "")
                            if nbs_text and nbs_text not in ("-", ""):
                                share_val = nbs_text
                        if natixis_residual_cash_col_index != -1 and len(cells) > natixis_residual_cash_col_index:
                            rc_text = cells[natixis_residual_cash_col_index].get_text(strip=True).replace(",", "")
                            if rc_text and rc_text not in ("-", ""):
                                cash_val = rc_text
                        if natixis_final_price_col_index != -1 and len(cells) > natixis_final_price_col_index:
                            fp_text = cells[natixis_final_price_col_index].get_text(strip=True).replace(",", "")
                            try:
                                natixis_closing_price = float(fp_text)
                            except (ValueError, TypeError):
                                pass

                    if clean_date and isin_val and len(isin_val) > 5:
                        extracted_data.append({
                            "ISIN": isin_val,
                            "Event Type": event_val,
                            "Value Day": clean_date,
                            "Obs Day": clean_obs_date,
                            "Value": display_val,
                            "Source": source_name,
                            "Denom": denom_val,
                            "Strike": "0",
                            "Shares": share_val,
                            "FractionalCash": cash_val,
                            "_ClosingPrice": natixis_closing_price,
                        })

    return extracted_data


def extract_data_from_text(text, source_type="Text"):
    """備用解析函式：對純文字用 regex 暴力掃描 ISIN 碼與日期"""
    extracted_items = []
    isins = re.findall(ISIN_PATTERN, text)
    dates = re.findall(DATE_PATTERN_TEXT, text)
    primary_date = dates[0] if dates else datetime.date.today().strftime("%Y-%m-%d")

    for isin in list(set(isins)):
        extracted_items.append({
            "ISIN": isin,
            "Value Day": primary_date,
            "Obs Day": "",
            "Source": source_type,
            "Denom": "10000",
            "Strike": "0",
            "Shares": "0",
            "FractionalCash": "0"
        })

    return extracted_items


def process_encrypted_excel(file_path, password):
    """在記憶體中解密加密的 Excel 檔案 (msoffcrypto)，不落地儲存解密後的明文檔案"""
    decrypted_workbook = io.BytesIO()
    try:
        with open(file_path, "rb") as file:
            office_file = msoffcrypto.OfficeFile(file)
            if password:
                office_file.load_key(password=password)
                office_file.decrypt(decrypted_workbook)
                return pd.read_excel(decrypted_workbook, sheet_name=None, engine='openpyxl')
            else:
                return pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
    except:
        try:
            return pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
        except:
            return {}


def process_bbva_excel(df, sheet_name="Sheet1"):
    """BBVA 專屬 Excel 解析"""
    extracted_data = []
    try:
        df.columns = df.columns.str.strip()
        col_map = {str(c).lower(): c for c in df.columns}

        target_date_key = 'payment date'
        target_isin_key = 'isin'
        target_obs_key = 'fixing date'

        if target_date_key in col_map and target_isin_key in col_map:
            real_date_col = col_map[target_date_key]
            real_isin_col = col_map[target_isin_key]
            real_event_col = col_map.get('event type', None)
            real_obs_col = col_map.get(target_obs_key, None)
            real_ticker_col = col_map.get('worstof ticker', None)

            real_share_col = None
            for sc in ['nb shares / denom.', 'nb shares / dom.', 'nb shares / nom.', 'nb shares/denom.']:
                if sc in col_map:
                    real_share_col = col_map[sc]
                    break

            real_cash_col = None
            for cc in ['res. cash / denom.', 'res. cash / dom.', 'res. cash / nom.', 'res. cash/denom.']:
                if cc in col_map:
                    real_cash_col = col_map[cc]
                    break

            df[real_date_col] = pd.to_datetime(df[real_date_col], errors='coerce')
            if real_obs_col:
                df[real_obs_col] = pd.to_datetime(df[real_obs_col], errors='coerce')

            for _, row in df.iterrows():
                isin_val = str(row[real_isin_col]).strip()

                ts = row[real_date_col]
                pay_date_str = ""
                if pd.notna(ts):
                    pay_date_str = ts.strftime('%Y-%m-%d')

                obs_date_str = ""
                if real_obs_col:
                    ts_obs = row[real_obs_col]
                    if pd.notna(ts_obs):
                        obs_date_str = ts_obs.strftime('%Y-%m-%d')

                event_type = "BBVA Fixing/Payment"
                if real_event_col and pd.notna(row[real_event_col]):
                    event_type = str(row[real_event_col]).strip()

                display_value = "BBVA Excel Processed"

                if "final redemption" in event_type.lower():
                    if real_ticker_col:
                        ticker_val = str(row[real_ticker_col]).strip()
                        if ticker_val and ticker_val.lower() != 'nan':
                            display_value = f"換股標的:{ticker_val}"

                shares_val = "0"
                if real_share_col and pd.notna(row[real_share_col]):
                    v = str(row[real_share_col]).strip()
                    if v.lower() != 'nan':
                        shares_val = v

                cash_val = "0"
                if real_cash_col and pd.notna(row[real_cash_col]):
                    c = str(row[real_cash_col]).strip()
                    if c.lower() != 'nan':
                        cash_val = c

                if isin_val and len(isin_val) > 5 and isin_val.lower() != 'nan':
                    extracted_data.append({
                        "ISIN": isin_val,
                        "Event Type": event_type,
                        "Value Day": pay_date_str,
                        "Obs Day": obs_date_str,
                        "Value": display_value,
                        "Source": f"BBVA Excel ({sheet_name})",
                        "Denom": "10000",
                        "Strike": "0",
                        "Shares": shares_val,
                        "FractionalCash": cash_val
                    })
        else:
            return []

    except Exception as e:
        print(f"BBVA Excel 解析發生錯誤: {e}")
        return []

    return extracted_data


def process_pdf(file_path):
    """一般 PDF 解析：擷取全頁文字後交給 extract_data_from_text 用 regex 掃描"""
    full_text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"
        return extract_data_from_text(full_text, "PDF")
    except:
        return []


def parse_excel_dataframe(df, company_name=""):
    """通用 Excel 解析器，適用於 JPM / UBS / Nomura / 一般 Excel"""
    results = []

    def normalize_col_name(col):
        if isinstance(col, tuple):
            col = " ".join([str(x) for x in col if str(x).strip().lower() != "nan"])
        col = str(col).strip().lower()
        col = col.replace("\n", " ").replace("\r", " ")
        col = col.replace("_", "").replace("-", "").replace("/", "").replace("\\", "")
        col = col.replace("(", "").replace(")", "").replace(".", "").replace(":", "").replace("*", "")
        col = re.sub(r"\s+", "", col)
        return col

    def safe_str(val):
        if pd.isna(val):
            return ""
        s = str(val).strip()
        if s.lower() == "nan":
            return ""
        return s

    def smart_date(raw_val):
        raw_val = safe_str(raw_val)
        if not raw_val:
            return ""
        use_us = any(x in str(company_name).upper() for x in ["JPM", "JPMORGAN", "MORGAN", "UBS", "NOMURA", "CHASE"])
        return convert_financial_date(raw_val, is_us_format=use_us)

    def parse_numeric_value(raw_val):
        text = safe_str(raw_val).replace(",", "")
        if not text:
            return None
        match = re.search(r"-?\d+(?:\.\d+)?", text)
        if not match:
            return None
        try:
            return float(match.group(0))
        except (ValueError, TypeError):
            return None

    def format_numeric_value(num):
        if num is None:
            return ""
        if float(num).is_integer():
            return str(int(num))
        return f"{num:.10f}".rstrip("0").rstrip(".")

    def get_first_nonempty_value(row, col_map, aliases, is_date=False):
        for alias in aliases:
            if alias in col_map:
                raw = row.get(col_map[alias], "")
                if is_date:
                    parsed = smart_date(raw)
                    if parsed:
                        return parsed
                else:
                    s = safe_str(raw)
                    if s:
                        return s
        return ""

    def extract_isin_from_row(row, col_map):
        isin_aliases = [
            "isin", "productisin", "noteisin", "securityisin",
            "securityid", "instrumentisin", "productid", "isincode"
        ]
        for alias in isin_aliases:
            if alias in col_map:
                val = safe_str(row.get(col_map[alias], ""))
                m = re.search(ISIN_PATTERN, val)
                if m:
                    return m.group(0)
        row_text = " ".join([safe_str(v) for v in row.values])
        m = re.search(ISIN_PATTERN, row_text)
        return m.group(0) if m else ""

    def detect_header_row(input_df, max_scan_rows=20):
        current_cols = [normalize_col_name(c) for c in input_df.columns]
        strong_header_keys = {
            "isin", "productisin", "noteisin", "securityisin",
            "eventtype", "paymentdate", "settlementdate", "observationdate",
            "eventdate", "evaluationdate", "valuationdate"
        }
        if any(c in strong_header_keys for c in current_cols):
            return input_df.copy()

        scan_n = min(max_scan_rows, len(input_df))
        for i in range(scan_n):
            row_vals = [normalize_col_name(v) for v in input_df.iloc[i].tolist()]
            hit_isin = any(v in ["isin", "productisin", "noteisin", "securityisin", "isincode"] for v in row_vals)
            hit_other = any(v in [
                "eventtype", "event", "eventstatus", "paymentdate", "settlementdate",
                "cashsettlementdate", "observationdate", "eventdate",
                "evaluationdate", "valuationdate", "fixingdate"
            ] for v in row_vals)
            if hit_isin and hit_other:
                new_df = input_df.copy()
                new_df.columns = [str(v).strip() for v in input_df.iloc[i].tolist()]
                new_df = new_df.iloc[i + 1:].reset_index(drop=True)
                return new_df

        return input_df.copy()

    def build_col_map(columns):
        col_map = {}
        for c in columns:
            norm = normalize_col_name(c)
            if norm and norm not in col_map:
                col_map[norm] = c
        return col_map

    try:
        if isinstance(df.columns, pd.MultiIndex):
            df = df.copy()
            df.columns = [
                " ".join([str(x) for x in col if str(x).strip().lower() != "nan"]).strip()
                for col in df.columns
            ]
    except Exception:
        pass

    df = detect_header_row(df)
    if df.empty:
        return results

    col_map = build_col_map(df.columns)
    df_str = df.astype(str)

    event_aliases = [
        "eventtype", "event", "eventstatus", "notificationtype",
        "redemptiontype", "coupontype", "paymenttype", "eventdescription",
        "description", "corporateactiontype", "actiontype"
    ]
    value_day_aliases = [
        "paymentdate", "valuedate", "settlementdate", "cashsettlementdate",
        "maturitydate", "redemptiondate", "couponpaymentdate",
        "scheduledpaymentdate", "actualpaymentdate"
    ]
    obs_day_aliases = [
        "obsdate", "observationdate", "eventdate", "evaluationdate",
        "valuationdate", "fixingdate", "finalvaluationdate",
        "pricingdate", "determinationdate", "observationday"
    ]
    denom_aliases = [
        "denomination", "denom", "nominal", "notional", "notionalamount",
        "principalamount", "faceamount", "specifieddenomination"
    ]
    strike_aliases = [
        "strikeprice", "strike", "initialprice", "referenceprice",
        "exerciseprice", "finalstrike", "strikelevel"
    ]
    shares_aliases = [
        "shares", "shareamount", "sharesperdenom", "sharesperdenomination",
        "numberofshares", "numberofsharespernote", "numberofsharesperdenom",
        "totalsharedeliver", "totalsharedelivery", "redeemedshareperdenomination",
        "quantityofstockdenom", "quantityofstockperdenom", "stockquantity",
        "deliveredshares", "sharestobedelivered", "sharedeliveryamount",
        "redemptionshares", "numberofphysicalshares", "nbsharesdenom",
        "nbsharesnom", "deliverysharepernote"
    ]
    cash_aliases = [
        "fractionalcash", "residualcash", "cashamount", "cashsettlementamount",
        "redemptionamount", "redemptionamountperdenom", "rescashdenom",
        "rescashperdenom", "cashpayment", "couponamount", "totalamountpayable",
        "cashcomponent", "fractionalentitlementcash", "residualamount",
        "residualcashamount", "cashinlieu", "netcashamount",
        "couponamountperdenom", "residualamountpernote",
        "redemptioncashperdenomination", "redemptionamtdenom"
    ]
    underlying_aliases = [
        "underlyingticker", "underlying", "redeemedshareticker",
        "redemptionshareticker", "worstofticker", "worstperf",
        "shareelement", "sharestobedelivered", "stockredemptionsedolisin",
        "stockredemption", "physicaldeliveryasset", "deliveredasset",
        "underlyingsecurity", "underlyingname", "ric"
    ]
    traded_notes_aliases = [
        "tradednumberofnotes", "numberoftradednotes", "tradednotes",
        "tradednumber", "numberofnotes"
    ]

    for _, row in df_str.iterrows():
        isin_code = extract_isin_from_row(row, col_map)
        if not isin_code:
            continue

        row_text = " ".join([safe_str(v) for v in row.values])

        event_val = get_first_nonempty_value(row, col_map, event_aliases, is_date=False)
        date_val = get_first_nonempty_value(row, col_map, value_day_aliases, is_date=True)
        obs_val = get_first_nonempty_value(row, col_map, obs_day_aliases, is_date=True)

        company_upper = str(company_name).upper()
        if any(x in company_upper for x in ["JPM", "JPMORGAN", "CHASE"]):
            settle_val = get_first_nonempty_value(row, col_map, ["settledate"], is_date=True)
            if settle_val:
                date_val = settle_val

        if not date_val:
            found_date = re.search(DATE_PATTERN_TEXT, row_text)
            if found_date:
                date_val = smart_date(found_date.group(0))

        denom_val = get_first_nonempty_value(row, col_map, denom_aliases, is_date=False) or "10000"
        strike_val = get_first_nonempty_value(row, col_map, strike_aliases, is_date=False) or "0"
        shares_val = get_first_nonempty_value(row, col_map, shares_aliases, is_date=False) or "0"
        cash_val = get_first_nonempty_value(row, col_map, cash_aliases, is_date=False) or "0"
        underlying_val = get_first_nonempty_value(row, col_map, underlying_aliases, is_date=False)

        if "UBS" in company_upper:
            delivery_share_per_note = get_first_nonempty_value(row, col_map, ["deliverysharepernote"], is_date=False)
            if delivery_share_per_note:
                shares_val = delivery_share_per_note

            residual_amount = get_first_nonempty_value(row, col_map, ["residualamount"], is_date=False)
            traded_number_of_notes = get_first_nonempty_value(row, col_map, traded_notes_aliases, is_date=False)
            residual_num = parse_numeric_value(residual_amount)
            traded_notes_num = parse_numeric_value(traded_number_of_notes)
            if residual_num is not None and traded_notes_num not in (None, 0):
                cash_val = format_numeric_value(residual_num / traded_notes_num)

        denom_val = denom_val.replace(",", "") if denom_val else "10000"
        strike_val = strike_val.replace(",", "") if strike_val else "0"
        shares_val = shares_val.replace(",", "") if shares_val else "0"
        cash_val = cash_val.replace(",", "") if cash_val else "0"

        row_upper = row_text.upper()
        event_upper = str(event_val).upper()

        if not event_val:
            if "COUPON" in row_upper:
                event_val = "Coupon"
            elif "EARLY REDEMPTION" in row_upper:
                event_val = "Early Redemption"
            elif "REDEMPTION" in row_upper:
                event_val = "Redemption"
            elif "MATURITY" in row_upper or "EXPIRY" in row_upper:
                event_val = "Expiry"
            elif "EXPIRED" in row_upper:
                event_val = "Expired"
            elif "FIXING" in row_upper:
                event_val = "Fixing"

        if "UBS" in company_upper:
            if not event_val:
                if "CLIENT NOTICE" in row_upper:
                    event_val = "Client Notice"
                elif "COUPON" in row_upper:
                    event_val = "Coupon"
                elif "REDEMPTION" in row_upper:
                    event_val = "Redemption"
            if str(event_val).strip().upper() == "COUPON":
                cash_val = "0"

        if any(x in company_upper for x in ["JPM", "JPMORGAN", "CHASE"]):
            if not event_val:
                if "SECURITIZED TRADE" in row_upper:
                    event_val = "JPMorgan Securitized Trade"
                elif "FIXING" in row_upper:
                    event_val = "Fixing Notification"
                elif "COUPON" in row_upper:
                    event_val = "Coupon"

        event_upper = str(event_val).upper()
        final_row_value = "See Excel"

        if underlying_val:
            is_early_redemption = "EARLY REDEMPTION" in event_upper
            has_share_quantity = parse_numeric_value(shares_val) not in (None, 0)
            has_delivery_indicator = (
                "REDEMPTION IN SHARES" in event_upper
                or "PHYSICAL" in event_upper
                or "DELIVER" in event_upper
                or "PHYSICAL DELIVERY" in row_upper
                or "DELIVERED" in row_upper
            )
            is_share_delivery_event = (
                "EXPIRY" in event_upper
                or "EXPIRED" in event_upper
                or "MATURITY" in event_upper
                or "FINAL REDEMPTION" in event_upper
                or has_delivery_indicator
            )
            # Natixis 的 Final Redemption 只有在明確帶出交割訊號或股數時才視為換股。
            if (
                "NATIXIS" in company_upper
                and "FINAL REDEMPTION" in event_upper
                and not has_share_quantity
                and not has_delivery_indicator
            ):
                is_share_delivery_event = False
            if is_share_delivery_event and not is_early_redemption:
                final_row_value = f"換股標的:{underlying_val}"

        if final_row_value == "See Excel":
            if ("WORSTOF" in row_upper or "DELIVERED" in row_upper or "PHYSICAL DELIVERY" in row_upper):
                if "EARLY REDEMPTION" not in event_upper:
                    for alias in underlying_aliases:
                        if alias in col_map:
                            tmp = safe_str(row.get(col_map[alias], ""))
                            if tmp:
                                final_row_value = f"換股標的:{tmp}"
                                break

        if final_row_value == "See Excel":
            compact_text = re.sub(r"\s+", " ", row_text).strip()
            if compact_text:
                final_row_value = compact_text[:200]

        results.append({
            "ISIN": isin_code,
            "Event Type": event_val if event_val else "Unknown",
            "Value Day": date_val if date_val else "Check File",
            "Obs Day": obs_val,
            "Value": final_row_value,
            "Source": "Excel",
            "Denom": denom_val,
            "Strike": strike_val,
            "Shares": shares_val,
            "FractionalCash": cash_val
        })

    return results


# ==========================================
# === HSBC OBS Date 查詢輔助函式 (Massive API) ===
# ==========================================

def _extract_hsbc_closing_table(html_body):
    """
    從 HSBC 換股贖回信件 HTML 底部的 ISIN/FX/strike/Closing 表格解析收盤價。
    回傳 dict，key = note ISIN (str)，value = closing price (float)。
    """
    closing_map = {}
    if not html_body:
        return closing_map
    soup = BeautifulSoup(html_body, "html.parser")
    for table in soup.find_all("table"):
        all_rows = table.find_all("tr")
        header_norm = None
        isin_ci = closing_ci = -1
        for tr in all_rows:
            cells = [td.get_text(strip=True) for td in tr.find_all(["th", "td"])]
            norm  = [c.upper() for c in cells]
            if "ISIN" in norm and "CLOSING" in norm:
                header_norm = norm
                isin_ci    = norm.index("ISIN")
                closing_ci = norm.index("CLOSING")
                break
        if header_norm is None:
            continue
        found_header = False
        for tr in all_rows:
            cells = [td.get_text(strip=True) for td in tr.find_all(["th", "td"])]
            norm  = [c.upper() for c in cells]
            if "ISIN" in norm and "CLOSING" in norm:
                found_header = True
                continue
            if not found_header:
                continue
            if len(cells) > max(isin_ci, closing_ci):
                raw_isin    = cells[isin_ci].strip()
                raw_closing = cells[closing_ci].strip()
                if not re.fullmatch(r"[A-Z]{2}[A-Z0-9]{9}\d", raw_isin):
                    continue
                try:
                    closing_map[raw_isin] = float(raw_closing.replace(",", ""))
                except ValueError:
                    pass
    return closing_map


def _extract_closing_from_text_backup(text, target_isin):
    """
    純文字備援：在包含 target_isin 的行中找最後一個浮點數作為 Closing。
    HSBC 格式通常為：XS3165690648  1  297.9125  284.64
    """
    for line in text.splitlines():
        if target_isin in line:
            nums = re.findall(r"[\d,]+\.\d+", line)
            if len(nums) >= 2:
                try:
                    return float(nums[-1].replace(",", ""))
                except ValueError:
                    pass
    return None


def _ticker_to_massive_format(ticker_str):
    """Bloomberg ticker 轉 Massive API 格式：'TSM UN' → 'TSM'"""
    if not ticker_str:
        return ""
    return ticker_str.strip().split()[0].upper()


def _fetch_daily_bars_massive(api_key, ticker, date_from, date_to):
    """呼叫 Massive API 取得日 K 線，回傳 payload dict。"""
    url = MASSIVE_BASE_URL.format(ticker=ticker, date_from=date_from, date_to=date_to)
    params = {"adjusted": "true", "sort": "asc", "limit": 5000, "apiKey": api_key}
    resp = requests.get(url, params=params, timeout=30)
    resp.raise_for_status()
    payload = resp.json()
    if payload.get("status") not in {"OK", "DELAYED"}:
        raise RuntimeError(f"Massive API 非正常狀態 [{ticker}]: {payload.get('status')}")
    return payload


def _find_obs_date_by_closing(bars_results, target_closing, value_day_str):
    """
    在日 K 線中精準比對 closing price (round 到 2 位小數)。
    若多筆符合，取最接近且 <= value_day 的日期。
    回傳 YYYY-MM-DD str 或 None。
    """
    if not bars_results or target_closing is None:
        return None
    target_rounded = round(float(target_closing), 2)
    try:
        value_day = datetime.datetime.strptime(value_day_str, "%Y-%m-%d").date()
    except Exception:
        value_day = None
    candidates = []
    for bar in bars_results:
        bar_close = bar.get("c")
        bar_ts    = bar.get("t")
        if bar_close is None or bar_ts is None:
            continue
        if round(float(bar_close), 2) != target_rounded:
            continue
        try:
            bar_date = datetime.datetime.fromtimestamp(bar_ts / 1000).date()
        except Exception:
            continue
        candidates.append(bar_date)
    if not candidates:
        return None
    if value_day:
        candidates = [d for d in candidates if d <= value_day]
    if not candidates:
        return None
    return max(candidates).strftime("%Y-%m-%d")


def _lookup_obs_date_from_massive(note_isin, underlying_ticker, value_day, closing_price,
                                  api_key=MASSIVE_DEFAULT_API_KEY):
    """
    整合查詢：給定 ISIN / 換股 Ticker / Value Day / Closing Price，
    呼叫 Massive API 找出對應的 OBS Date。
    回傳 YYYY-MM-DD str 或 None（查無或異常）。
    """
    ticker = _ticker_to_massive_format(underlying_ticker)
    if not ticker or not value_day or closing_price is None:
        return None
    try:
        vd = datetime.datetime.strptime(value_day, "%Y-%m-%d").date()
    except Exception:
        return None
    date_from = (vd - datetime.timedelta(days=30)).strftime("%Y-%m-%d")
    date_to   = vd.strftime("%Y-%m-%d")
    try:
        print(f"  [Massive] 查詢 {ticker} [{date_from}~{date_to}] Closing={closing_price}")
        payload  = _fetch_daily_bars_massive(api_key, ticker, date_from, date_to)
        bars     = payload.get("results", [])
        obs_date = _find_obs_date_by_closing(bars, closing_price, value_day)
        if obs_date:
            print(f"  [Massive] ✓ OBS Date={obs_date}")
        else:
            print(f"  [Massive] ✗ 未找到符合 Closing={closing_price} 的日期")
        return obs_date
    except Exception as e:
        print(f"  [Massive] API 錯誤: {e}")
        return None


def step1_extract_from_outlook(start_date, end_date):
    """
    步驟 1 (ELN1)：連接 Outlook → 逐封信件解析 → 輸出 ELN_Report Excel 報表
    回傳輸出檔名，失敗時回傳 None
    """
    output_filename = build_report_output_path(start_date, end_date)
    print("\n=== 步驟 1：Outlook 信件擷取開始 ===")

    try:
        outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
        target_folder = get_target_outlook_folder(
            outlook_namespace=outlook,
            account_name=OUTLOOK_ACCOUNT,
            target_folder_name=TARGET_FOLDER_NAME
        )
        print(f"成功鎖定帳號：{OUTLOOK_ACCOUNT}")
        print(f"成功鎖定資料夾：{TARGET_FOLDER_NAME}")
    except Exception as e:
        print(f"連線 Outlook 錯誤或找不到帳號/資料夾：{e}")
        return None

    messages = get_outlook_messages(target_folder, start_date, end_date)

    final_data = []
    count_processed = 0

    for msg in messages:
        try:
            if msg.Class != 43:  # Class 43 = MailItem
                continue

            mail_date = msg.ReceivedTime.date()
            if mail_date < start_date or mail_date > end_date:
                continue

            count_processed += 1
            sender_name = msg.SenderName
            try:
                sender_email = msg.SenderEmailAddress
            except:
                sender_email = ""

            extracted_company = extract_company_domain(sender_email, sender_name)
            subject = msg.Subject
            received_time_str = msg.ReceivedTime.strftime("%Y-%m-%d")

            # 跳過 Upcoming events 相關的信件
            if "Upcoming events" in subject:
                print(f"[{count_processed}] 跳過：{extracted_company} | {subject[:50]} (Upcoming events)")
                continue

            print(f"[{count_processed}] {extracted_company} | {subject[:50]}")

            general_event_type = get_event_type(msg.Body)
            if general_event_type == "Unknown":
                general_event_type = get_event_type(subject)

            products_found = []

            # 防呆：偵測從未見過的新上手寄件人 (在任何 if/elif 條件觸發前先檢查)
            if not any(k in extracted_company for k in KNOWN_COMPANIES):
                print(f"\n{'!' * 60}")
                print(f"[防呆警示] 發現從未見過的新上手信！")
                print(f"  寄件人識別碼 : {extracted_company}")
                print(f"  寄件者 Email : {sender_email}")
                print(f"  信件主旨     : {subject}")
                print(f"  收信日期     : {received_time_str}")
                print(f"  → 請確認是否需要為此寄件人新增專屬解析邏輯！")
                print(f"{'!' * 60}\n")

            # === HTML 處理區 ===
            if "SGMARKETS" in extracted_company:
                # 只處理主旨為 "Events on your Structured Products" 的 SG 信件
                if "Events on your Structured Products" not in subject:
                    print(f"[SG 跳過] 主旨不符，略過此信: {subject}")
                    continue
                try:
                    sg_data = process_sg_markets_html(msg.HTMLBody)
                    if sg_data:
                        products_found.extend(sg_data)
                except:
                    pass

            elif "BARCLAYS" in extracted_company:
                try:
                    barc_data = process_generic_html_table(
                        msg.HTMLBody,
                        [r"Value\s*Date"],
                        "Barclays HTML",
                        is_us_format=False
                    )
                    if barc_data:
                        products_found.extend(barc_data)
                except:
                    pass

            elif "MORGAN" in extracted_company:
                # 跳過回覆信（主旨含 RE: 或 [EXTERNAL]），只處理原始通知信
                if re.search(r"^\s*RE\s*:", subject, re.IGNORECASE) or "[EXTERNAL]" in subject:
                    print(f"[MS 跳過] 回覆信，略過此信: {subject}")
                    continue
                try:
                    # Morgan Stanley 支援多種表格格式，包括 KO/Expiry Notification (Expiration - Shares)
                    ms_keywords = [r"Payment\s*Date", r"Cash\s*Settlement\s*Date", r"Settlement\s*Date"]
                    ms_data = process_generic_html_table(
                        msg.HTMLBody,
                        ms_keywords,
                        "Morgan Stanley HTML",
                        is_us_format=True
                    )
                    if ms_data:
                        products_found.extend(ms_data)
                except:
                    pass

            elif "HSBC" in extracted_company:
                try:
                    hsbc_html_data = process_hsbc_html(msg.HTMLBody, subject)
                    if hsbc_html_data:
                        products_found.extend(hsbc_html_data)
                    elif "redemption in shares" in subject.lower():
                        hsbc_text_data = process_hsbc_redemption_shares(msg.Body)
                        if hsbc_text_data:
                            # 抓 HTML 底部 Closing 表格，查詢 OBS Date via Massive API
                            hsbc_closing_map = _extract_hsbc_closing_table(msg.HTMLBody)
                            for prod in hsbc_text_data:
                                prod_isin  = prod.get("ISIN", "")
                                closing_px = hsbc_closing_map.get(prod_isin) or \
                                             _extract_closing_from_text_backup(msg.Body, prod_isin)
                                # Ticker 從 Value 欄取出：例如 "換股標的: TSM UN (US8740391003)"
                                raw_val = prod.get("Value", "")
                                tmatch  = re.search(r"換股標的:\s*([^\(]+)", raw_val)
                                und_ticker = tmatch.group(1).strip() if tmatch else ""
                                if closing_px and und_ticker:
                                    obs_found = _lookup_obs_date_from_massive(
                                        prod_isin, und_ticker,
                                        prod.get("Value Day", ""), closing_px
                                    )
                                    if obs_found:
                                        prod["Obs Day"] = obs_found
                            products_found.extend(hsbc_text_data)
                except:
                    pass

            elif "NATIXIS" in extracted_company:
                try:
                    natixis_keywords = [r"PAYMENT\s*DATE"]
                    natixis_data = process_generic_html_table(
                        msg.HTMLBody, natixis_keywords, "Natixis HTML", is_us_format=False
                    )

                    try:
                        nat_soup = BeautifulSoup(msg.HTMLBody, 'html.parser')
                        for table in nat_soup.find_all('table'):
                            all_trs = table.find_all('tr')
                            nat_header = None
                            isin_ci = redeem_ci = -1
                            for tr in all_trs:
                                cells = [re.sub(r'\s+', ' ', td.get_text(separator=' ', strip=True)).strip()
                                         for td in tr.find_all(['th', 'td'])]
                                norm = [c.upper() for c in cells]
                                has_isin = any(c == 'ISIN' for c in norm)
                                redeem_idx = next((ci for ci, c in enumerate(norm) if c == 'REDEMPTION'), -1)
                                if has_isin and redeem_idx != -1:
                                    nat_header = norm
                                    isin_ci = next(ci for ci, c in enumerate(norm) if c == 'ISIN')
                                    redeem_ci = redeem_idx
                                    continue
                                if nat_header is not None and isin_ci != -1 and redeem_ci != -1:
                                    isin_val = cells[isin_ci] if isin_ci < len(cells) else ""
                                    redeem_val = cells[redeem_ci] if redeem_ci < len(cells) else ""
                                    if isin_val and redeem_val:
                                        for item in natixis_data:
                                            if item.get("ISIN") == isin_val:
                                                item["Event Type"] = redeem_val
                            if nat_header is not None:
                                break
                    except:
                        pass

                    obs_date_from_subj = ""
                    val_match = re.search(
                        r"Valuation\s*Date\s*(\d{1,2}[- ]?[A-Za-z]{3}[- ]?\d{2,4})",
                        subject, re.IGNORECASE
                    )
                    if val_match:
                        obs_date_from_subj = convert_financial_date(val_match.group(1))

                    if natixis_data:
                        # OBS Date 查詢：Natixis FINAL 換股轉換 → Yahoo Finance 歷史收盤價比對
                        for item in natixis_data:
                            closing_px = item.pop("_ClosingPrice", None)
                            raw_val = item.get("Value", "")
                            tmatch = re.search(r"換股標的:\s*([^\(]+)", raw_val)
                            und_ticker = tmatch.group(1).strip() if tmatch else ""
                            if closing_px and und_ticker:
                                obs_found = _lookup_obs_date_from_yahoo(
                                    und_ticker,
                                    item.get("Value Day", ""),
                                    closing_px
                                )
                                if obs_found:
                                    item["Obs Day"] = obs_found
                        if obs_date_from_subj:
                            for item in natixis_data:
                                if not item.get("Obs Day"):
                                    item["Obs Day"] = obs_date_from_subj
                        products_found.extend(natixis_data)
                except:
                    pass

            # === 附件處理區 ===
            excel_password = None
            if extracted_company in PASSWORD_MAP:
                excel_password = PASSWORD_MAP[extracted_company]
            else:
                for company_key, pwd in PASSWORD_MAP.items():
                    if company_key in sender_name.upper():
                        excel_password = pwd
                        break

            if msg.Attachments.Count > 0:
                for att in msg.Attachments:
                    filename = att.FileName
                    temp_file_path = os.path.join(TEMP_PATH, filename)

                    if not filename.lower().endswith((".pdf", ".xls", ".xlsx", ".xlsm")):
                        continue

                    try:
                        att.SaveAsFile(temp_file_path)

                        if filename.lower().endswith(".pdf"):
                            if "BNP" in extracted_company or "BNPPARIBAS" in extracted_company:
                                products_found.extend(process_bnp_pdf(temp_file_path, subject))
                            else:
                                products_found.extend(process_pdf(temp_file_path))

                        elif filename.lower().endswith((".xls", ".xlsx", ".xlsm")):
                            dfs_dict = process_encrypted_excel(temp_file_path, excel_password)
                            for sheet_name, df in dfs_dict.items():
                                if not df.empty:
                                    sheet_data = []
                                    if "BBVA" in extracted_company:
                                        sheet_data = process_bbva_excel(df, sheet_name)
                                        if not sheet_data:
                                            sheet_data = parse_excel_dataframe(df, extracted_company)
                                    else:
                                        sheet_data = parse_excel_dataframe(df, extracted_company)

                                    for item in sheet_data:
                                        if "Source" not in item:
                                            item["Source"] = f"Excel ({sheet_name})"
                                    products_found.extend(sheet_data)

                    except Exception:
                        pass
                    finally:
                        try:
                            if os.path.exists(temp_file_path):
                                os.remove(temp_file_path)
                        except:
                            pass

            if not products_found:
                products_found.extend(extract_data_from_text(msg.Body, "Email Body"))

            if products_found:
                for prod in products_found:
                    final_event_type = prod.get("Event Type")
                    if not final_event_type or final_event_type in ["None", "Unknown"]:
                        final_event_type = general_event_type

                    final_data.append({
                        "Date": received_time_str,
                        "Company": extracted_company,
                        "Title": subject,
                        "Event Type": final_event_type,
                        "ISIN": prod.get("ISIN"),
                        "ObsDate": prod.get("Obs Day", ""),
                        "Value Day/Payment Date": prod.get("Value Day"),
                        "Value/Raw Data": prod.get("Value", "詳見附件"),
                        "Source File": prod.get("Source"),
                        "Denom": prod.get("Denom", "10000"),
                        "Strike": prod.get("Strike", "0"),
                        "Shares": prod.get("Shares", "0"),
                        "FractionalCash": prod.get("FractionalCash", "0"),
                    })
            else:
                final_data.append({
                    "Date": received_time_str,
                    "Company": extracted_company,
                    "Title": subject,
                    "Event Type": general_event_type,
                    "ISIN": "Not Found",
                    "ObsDate": "",
                    "Value Day/Payment Date": "",
                    "Value/Raw Data": "",
                    "Source File": "No Data",
                    "Denom": "10000",
                    "Strike": "0",
                    "Shares": "0",
                    "FractionalCash": "0",
                })

        except Exception as e:
            print(f"略過信件錯誤: {e}")

    if final_data:
        df_result = pd.DataFrame(final_data)
        cols = [
            "Date", "Company", "Title", "Event Type", "ISIN", "ObsDate",
            "Value Day/Payment Date", "Value/Raw Data", "Source File",
            "Denom", "Strike", "Shares", "FractionalCash"
        ]
        df_result = df_result[cols]

        # 防呆：彙整所有 Unknown Event Type 並一次印出警示
        unknown_rows = df_result[df_result["Event Type"].astype(str).str.strip().str.upper() == "UNKNOWN"]
        if not unknown_rows.empty:
            print(f"\n{'!' * 60}")
            print(f"[防呆警示] 有 {len(unknown_rows)} 筆信件的 Event Type 無法辨識！")
            print(f"  → 請確認是否需要在 EVENT_KEYWORDS 中新增對應關鍵字：")
            for _, r in unknown_rows.iterrows():
                print(f"  [{r['Date']}] {r['Company']} | {str(r['Title'])[:60]}")
            print(f"{'!' * 60}\n")

        df_result.to_excel(output_filename, index=False)
        print(f"\n[步驟1完成] 報表已生成: {output_filename}")
        print(f"共擷取 {len(df_result)} 筆資料")
        return output_filename
    else:
        print(f"\n[步驟1結果] 區間內無相符資料。")
        return None


# ==========================================
# === 部位狀態檢查共用函式 ===
# ==========================================

def build_position_process_df(df):
    """
    部位狀態檢查：以 ISIN + 投資人帳號 為群組維度，
    保留「成交種類」開頭為 '1'（委託申購）的紀錄，
    但若同一群組內同時存在開頭為 '9'（提前解約）的紀錄，
    則整組剔除並印出軌跡訊息。
    防呆：若缺少 '投資人帳號' 或 '成交種類' 欄位，自動降級處理。
    回傳乾淨的 process_df。
    """
    # ── 防呆：欄位存在性檢查 ──
    missing = [c for c in ['投資人帳號', '成交種類'] if c not in df.columns]
    if missing:
        print(f"[部位狀態檢查] 警告：缺少欄位 {missing}，略過部位狀態檢查，直接回傳原始資料。")
        return df.copy()

    # ── Step 1：僅保留「成交種類」開頭為 '1' 的委託申購紀錄 ──
    purchase_df = df[df['成交種類'].astype(str).str.strip().str.startswith('1')].copy()

    # ── Step 2：找出同一群組內含有「成交種類」開頭為 '9' 的 ISIN+帳號組合 ──
    terminated_mask = df['成交種類'].astype(str).str.strip().str.startswith('9')
    terminated_pairs = (
        df[terminated_mask]
        .groupby(['ISIN', '投資人帳號'])
        .size()
        .reset_index()[['ISIN', '投資人帳號']]
    )

    # ── Step 3：迭代印出被剔除的軌跡紀錄 ──
    if not terminated_pairs.empty:
        for _, row in terminated_pairs.iterrows():
            isin    = row['ISIN']
            account = row['投資人帳號']
            print(f"【軌跡紀錄】投資人帳號 {account} 於 ISIN {isin} 已有 9.提前解約 紀錄，該筆略過處理。")

    # ── Step 4：從委託申購清單中排除已解約的群組 ──
    if not terminated_pairs.empty:
        terminated_pairs['_exclude'] = True
        purchase_df = purchase_df.merge(
            terminated_pairs,
            on=['ISIN', '投資人帳號'],
            how='left'
        )
        purchase_df = purchase_df[purchase_df['_exclude'].isna()].drop(columns=['_exclude'])

    print(f"[部位狀態檢查] 完成：共保留 {len(purchase_df)} 筆有效委託申購明細。")
    return purchase_df.reset_index(drop=True)


# ==========================================
# === 步驟 2：ELN2 — 合併統計表 ===
# ==========================================

def step2_merge_with_stats(start_date, end_date, report_file):
    """
    步驟 2 (ELN2)：讀取 ELN_Report → 合併 ELN統計表.xlsx → 輸出 ELN_Final_Output
    回傳輸出檔名，失敗時回傳 None
    """
    stats_file = resolve_runtime_path("ELN統計表.xlsx")
    output_filename = build_final_output_path(start_date, end_date)

    print("\n=== 步驟 2：合併統計表開始 ===")

    if not os.path.exists(report_file):
        print(f"錯誤：找不到步驟1的報表 '{report_file}'")
        return None

    if not os.path.exists(stats_file):
        print(f"錯誤：找不到 '{stats_file}'，請確認該檔案存在。")
        return None

    try:
        df_report = pd.read_excel(report_file)
        df_stats = pd.read_excel(stats_file)

        target_columns = ['ISIN', '下單日期', '投資人帳號', "分公司", '姓名', "成交種類", '商品名稱', '本日贖回單位數', '幣別']
        missing_cols = [col for col in target_columns if col not in df_stats.columns]

        if missing_cols:
            print(f"警告：'{stats_file}' 缺少以下欄位: {missing_cols}")
            return None

        df_stats_subset = df_stats[target_columns]

        df_report['ISIN'] = df_report['ISIN'].astype(str).str.strip()
        df_stats_subset = df_stats_subset.copy()
        df_stats_subset['ISIN'] = df_stats_subset['ISIN'].astype(str).str.strip()

        df_final = pd.merge(df_report, df_stats_subset, on='ISIN', how='left')
        df_final.to_excel(output_filename, index=False)

        print(f"[步驟2完成] 輸出檔案: {output_filename}")
        print(f"總資料筆數: {len(df_final)} 筆")
        return output_filename

    except Exception as e:
        print(f"步驟2發生錯誤: {e}")
        return None


# ==========================================
# === 步驟 3：ELN3 — 產生 Word 通知書 ===
# ==========================================

def step3_generate_documents(start_date, end_date, final_file):
    """
    步驟 3 (ELN3)：讀取 ELN_Final_Output → 分類事件類型 → 產生 Word 通知書
    """
    print("\n=== 步驟 3：產生 Word 通知書開始 ===")

    if not os.path.exists(final_file):
        print(f"錯誤：找不到步驟2的輸出檔 '{final_file}'")
        return

    try:
        df = pd.read_excel(final_file)
        df.columns = df.columns.str.strip()
    except Exception as e:
        print(f"Excel 讀取失敗: {e}")
        return

    if 'Event Type' not in df.columns:
        print("錯誤：Excel 中找不到 'Event Type' 欄位")
        return

    has_raw_data_col = 'Value/Raw Data' in df.columns
    if not has_raw_data_col:
        print("注意：Excel 中找不到 'Value/Raw Data' 欄位，將無法精準判斷換股標的。")

    # 關鍵字清單 (全部小寫比對)
    coupon_keywords = [
        'coupon', 'interest payment', 'coupon payment'
    ]
    early_keywords = [
        'early redemption', 'early termination', 'autocall',
        'coupon + early redemption', 'knockedout', 'early redemptions', 'coupon+redemption'
    ]
    conversion_specific_keywords = [
        'redemption in shares', 'expiration - shares', 'expiration-shares'
    ]
    ambiguous_keywords = [
        'maturity', 'final redemption', 'expiry', 'redemptions',
        'redemption at maturity', 'coupon + final redemption', 'expiration', 'redemption'
    ]

    def classify_row(row):
        event_type_lower = str(row.get('Event Type', '')).strip().lower()
        raw_data = str(row.get('Value/Raw Data', '')).strip()
        company_upper = str(row.get('Company', '')).upper()

        if event_type_lower in coupon_keywords:
            return 'coupon'
        if "UBS" in company_upper and event_type_lower == 'expired':
            return 'conversion'
        if event_type_lower in early_keywords:
            return 'early'
        if event_type_lower in conversion_specific_keywords:
            return 'conversion'
        if event_type_lower in ambiguous_keywords:
            if not has_raw_data_col or raw_data in ['', 'nan', '0']:
                return 'maturity'
            raw_lower = raw_data.lower()
            if "換股標的" in raw_data or "equity" in raw_lower or "share" in raw_lower:
                return 'conversion'
            return 'maturity'
        return None

    process_df = build_position_process_df(df)
    if process_df.empty and '成交種類' not in df.columns:
        process_df = df.drop_duplicates(subset=['ISIN', '姓名']).copy()

    print(f"共發現 {len(process_df)} 筆有效資料，開始分類並製作文件...")

    templates = {
        'early': resolve_runtime_path("境外結構型商品重要訊息提前.docx"),
        'maturity': resolve_runtime_path("境外結構型商品重要訊息到期.docx"),
        'conversion': resolve_runtime_path("境外結構型商品重要訊息轉換.docx")
    }

    templates['early'] = resolve_runtime_path("境外結構型商品重要訊息提前.docx")
    templates['maturity'] = resolve_runtime_path("境外結構型商品重要訊息到期.docx")
    templates['conversion'] = resolve_runtime_path("境外結構型商品重要訊息轉換.docx")

    for key, path in templates.items():
        if not os.path.exists(path):
            print(f"警告：找不到範本: {path}，該類型的文件將無法產出。")

    stats = {'early': 0, 'maturity': 0, 'conversion': 0, 'skip': 0}

    def fmt_date(d):
        try:
            return pd.to_datetime(d).strftime('%Y/%m/%d')
        except:
            return str(d)

    def replace_text(doc_obj, replacements):
        for paragraph in doc_obj.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, str(value))

    start_str = str(start_date)
    end_str = str(end_date)

    for index, row in process_df.iterrows():
        doc_type = classify_row(row)
        if doc_type == 'coupon':
            stats['skip'] += 1
            continue
        if doc_type is None:
            event_val = str(row.get('Event Type', '')).strip()
            print(f"[防呆警示] 無法分類 Event Type：「{event_val}」"
                  f" | ISIN: {row.get('ISIN', '')} | 公司: {row.get('Company', '')}"
                  f" → 請確認 classify_row() 或 early/ambiguous 關鍵字清單是否需要更新！")
            stats['skip'] += 1
            continue

        template_name = templates.get(doc_type)
        if not os.path.exists(template_name):
            print(f"跳過 ISIN {row.get('ISIN')}，因缺少範本 {template_name}")
            continue

        type_folder_name = {'early': '提前出場', 'maturity': '到期結算', 'conversion': '轉換股票'}.get(doc_type)
        output_folder = ensure_directory(WORD_OUTPUT_DIR)

        client_name = row.get('姓名', '客戶')
        isin = row.get('ISIN', '')
        trade_date = fmt_date(row.get('下單日期', ''))
        product_name = row.get('商品名稱', '')
        valuation_date = fmt_date(row.get('ObsDate', row.get('Date', '')))
        payment_date = fmt_date(row.get('Value Day/Payment Date', row.get('Payment Date', '')))
        raw_data_value = row.get('Value/Raw Data', '')

        replacements = {
            "{{姓名}}": client_name,
            "{{下單日期}}": trade_date,
            "{{ISIN}}": isin,
            "{{商品名稱}}": product_name,
            "{{ObsDate}}": valuation_date,
            "{{Date}}": valuation_date,
            "{{Value Day/Payment Date}}": payment_date,
            "{{PaymentDate}}": payment_date,
            "{{Value/Raw Data}}": raw_data_value,
            "{{幾年幾月}}": datetime.datetime.now().strftime("%Y年.%m月")
        }

        save_name = f"{type_folder_name}通知書_{str(isin)}_{str(client_name)}.docx".replace('/', '_')
        save_path = os.path.join(output_folder, save_name)

        try:
            doc = Document(template_name)
            replace_text(doc, replacements)
            doc.save(save_path)
            print(f"[{type_folder_name}] 已產出: {save_name}")
            stats[doc_type] += 1
        except Exception as e:
            print(f"產出失敗 ({save_name}): {e}")

    print("\n" + "=" * 50)
    print("步驟3執行完成！統計結果：")
    print(f" - 提前出場: {stats['early']} 筆")
    print(f" - 到期結算: {stats['maturity']} 筆")
    print(f" - 轉換股票: {stats['conversion']} 筆")
    print(f" - 未分類/跳過: {stats['skip']} 筆")
    print("=" * 50)


# ==========================================
# === 步驟 4：ELN(4) — 產生 Outlook 草稿 ===
# ==========================================

def normalize_company_key(company_text):
    """將公司名稱文字對應到統一的內部識別碼（SG/HSBC/BNP/JPM/MORGAN/BARCLAYS/BBVA/UBS/NOMURA/NATIXIS）。"""
    s = str(company_text or "").upper().strip()
    mapping = {
        "SGMARKETS": "SG",
        "SOCIETE GENERALE": "SG",
        "SOCIETE": "SG",
        "HSBC": "HSBC",
        "BNP": "BNP",
        "BNPPARIBAS": "BNP",
        "J.P. MORGAN": "JPM",
        "JPMORGAN": "JPM",
        "JPM": "JPM",
        "CHASE": "JPM",
        "MORGAN STANLEY": "MORGAN",
        "MORGAN": "MORGAN",
        "BARCLAYS": "BARCLAYS",
        "BBVA": "BBVA",
        "UBS": "UBS",
        "NOMURA": "NOMURA",
        "NATIXIS": "NATIXIS",
    }
    for k, v in mapping.items():
        if k in s:
            return v
    return s


def is_company_match(row_company, mail_company, sender_name="", subject=""):
    """將 Excel 列的 Company 欄與信件發件人資訊對比，支援同義詞群組。"""
    a = normalize_company_key(row_company)
    candidates = " ".join([str(mail_company or ""), str(sender_name or ""), str(subject or "")]).upper()
    b = normalize_company_key(candidates)
    if a == b:
        return True
    synonym_groups = [
        {"SG", "SGMARKETS"},
        {"BNP", "BNPPARIBAS"},
        {"JPM", "JPMORGAN", "CHASE"},
        {"MORGAN", "MORGAN STANLEY"},
    ]
    for grp in synonym_groups:
        if a in grp and b in grp:
            return True
    return False


def classify_row(row, has_raw_data_col):
    """核心判斷邏輯函式 (精準比對換股字眼)，回傳 'conversion'/'early'/'maturity'/None。"""
    early_keywords = [
        'early redemption', 'early termination', 'autocall',
        'knockedout', 'early redemptions', 'coupon+redemption',
        'ko/expiry', 'early', 'coupon + early redemption'
    ]
    conversion_specific_keywords = [
        'redemption in shares', 'expiration - shares', 'expiration-shares'
    ]
    ambiguous_keywords = [
        'maturity', 'final redemption', 'expiry', 'redemptions',
        'redemption at maturity', 'coupon + final redemption', 'expiration',
        'redemption payment', 'redemption'
    ]
    event_type = str(row.get('Event Type', '')).strip()
    event_type_lower = event_type.lower()
    raw_data = str(row.get('Value/Raw Data', '')).strip()
    raw_lower = raw_data.lower()
    company = str(row.get('Company', '')).upper()
    if 'UBS' in company and event_type_lower == 'expired':
        return 'conversion'
    for kw in conversion_specific_keywords:
        if kw in event_type_lower:
            return 'conversion'
    if "換股標的" in raw_data or "equity" in raw_lower or "share" in raw_lower:
        return 'conversion'
    for kw in early_keywords:
        if kw in event_type_lower:
            return 'early'
    if 'NATIXIS' in company and 'coupon' in event_type_lower:
        return None
    for kw in ambiguous_keywords:
        if kw in event_type_lower:
            return 'maturity'
    if (not event_type_lower or event_type_lower == 'nan' or event_type_lower == 'unknown') and has_raw_data_col:
        if "換股標的" in raw_data:
            return 'conversion'
    return None


def create_html_table(df):
    """將 Pandas DataFrame 轉換成 Outlook 支援的 HTML 表格。"""
    if df is None or df.empty:
        return ""
    table_html = '<table style="border-collapse: collapse; font-family: \'Microsoft JhengHei\', sans-serif; font-size: 10pt; color: black;">'
    table_html += '<tr style="background-color: #f2f2f2;">'
    for col in df.columns:
        table_html += f'<th style="border: 1px solid black; padding: 4px 8px; text-align: center;">{html.escape(str(col))}</th>'
    table_html += '</tr>'
    for _, row in df.iterrows():
        table_html += '<tr>'
        for val in row:
            display_val = "" if pd.isna(val) else str(val)
            table_html += f'<td style="border: 1px solid black; padding: 4px 8px; text-align: center;">{html.escape(display_val)}</td>'
        table_html += '</tr>'
    table_html += '</table>'
    return table_html


def normalize_event_type_to_category(event_type_text, raw_data_text=""):
    """將 Event Type 文字對應到 'conversion'/'early'/'maturity'，無法辨識則回傳 None。"""
    text = f"{event_type_text or ''} {raw_data_text or ''}".strip().lower()
    if not text or text == 'nan':
        return None
    conversion_keywords = [
        'redemption in shares', 'expiration - shares', 'share settlement',
        'physical settlement', 'settlement in shares', 'deliver shares',
        'share delivery', 'equity'
    ]
    early_keywords = [
        'early redemption', 'early termination', 'autocall', 'knockout',
        'knocked out', 'knockedout', 'coupon + early redemption',
        'coupon+redemption', 'ko/expiry'
    ]
    maturity_keywords = [
        'expiration', 'maturity', 'final redemption', 'redemption at maturity',
        'coupon + final redemption', 'coupon+final redemption',
        'redemption payment'
    ]
    for kw in conversion_keywords:
        if kw in text:
            return 'conversion'
    for kw in early_keywords:
        if kw in text:
            return 'early'
    for kw in maturity_keywords:
        if kw in text:
            return 'maturity'
    return None


def content_text_matches_category(text, category):
    """判斷一段自由文字是否符合指定分類（conversion / early / maturity）。"""
    if not text or not category:
        return False
    return normalize_event_type_to_category(text) == category


def normalize_event_type_label(value):
    """將 Event Type 字串統一化：壓縮空白、轉小寫、全形加號轉半形。"""
    text = re.sub(r'\s+', ' ', str(value or '').strip()).lower()
    text = text.replace('＋', '+')
    return text


def text_matches_target_event_type(text, target_event_type=None, target_category=None):
    """判斷文字是否符合目標 Event Type（精確子字串比對）或目標分類（category）。"""
    if not text:
        return False
    normalized_text = normalize_event_type_label(text)
    normalized_target = normalize_event_type_label(target_event_type)
    if normalized_target:
        if normalized_target in normalized_text:
            return True
        alt_target = normalized_target.replace(' - ', '-').replace(' + ', '+')
        alt_text = normalized_text.replace(' - ', '-').replace(' + ', '+')
        if alt_target and alt_target in alt_text:
            return True
    if target_category:
        return normalize_event_type_to_category(text) == target_category
    return False


def row_matches_isin_and_event(row_text, target_isin, target_event_type=None, target_category=None, allow_category_fallback=True):
    """判斷一列合併文字是否同時包含目標 ISIN 並符合目標 Event Type / 分類。"""
    if not cell_contains_target(row_text, target_isin):
        return False
    if target_event_type:
        if text_matches_target_event_type(row_text, target_event_type=target_event_type, target_category=target_category):
            return True
        return allow_category_fallback and bool(target_category) and content_text_matches_category(row_text, target_category)
    if target_category:
        return content_text_matches_category(row_text, target_category)
    return True


def normalize_text_for_match(value):
    """將儲存格內容統一化：壓縮連續空白、去首尾空白並轉大寫，供 ISIN 比對使用。"""
    return re.sub(r'\s+', ' ', str(value or '')).strip().upper()


def cell_contains_target(value, target):
    """判斷儲存格內容（經標準化）是否包含目標字串（大小寫不敏感）。"""
    if not target:
        return False
    return normalize_text_for_match(target) in normalize_text_for_match(value)


def row_values_joined(values):
    """將一列的所有非空儲存格內容以 ' | ' 連接為單一字串。"""
    return ' | '.join([str(v) for v in values if v not in (None, '')])


def find_header_row_above(ws, matched_row, max_col):
    """在 matched_row 上方 8 列範圍內搜尋最可能的表頭列，回傳評分最高的列號。"""
    header_keywords = ['isin', 'event', 'event type', 'product', 'trade', 'payment', 'redemption', 'coupon', 'shares']
    start_row = max(1, matched_row - 8)
    best_row = 1
    best_score = -1
    for r in range(start_row, matched_row + 1):
        score = 0
        for c in range(1, max_col + 1):
            val = ws.cell(r, c).value
            sval = str(val or '').strip().lower()
            if not sval:
                continue
            if any(k in sval for k in header_keywords):
                score += 3
            if sval == 'isin':
                score += 5
            if sval == 'event type':
                score += 5
        if score > best_score:
            best_score = score
            best_row = r
    return best_row


def html_escape_keep_breaks(value):
    """對儲存格內容進行 HTML 跳脫，並將換行符 \\n 保留為 <br>。"""
    s = "" if value is None else str(value)
    return html.escape(s).replace("\n", "<br>")


def excel_color_to_css(rgb):
    """將 openpyxl 讀到的 ARGB 8位十六進位色碼轉換為 CSS #RRGGBB格式。"""
    if not rgb:
        return ""
    rgb = str(rgb)
    if len(rgb) == 8:
        rgb = rgb[2:]
    if len(rgb) == 6 and rgb.upper() != "000000":
        return f"#{rgb}"
    return ""


def worksheet_selected_rows_to_html(ws, selected_rows, header_row=None):
    """將工作表中指定列（含可選表頭列）轉換為 inline-CSS HTML 表格。"""
    if not selected_rows:
        return ''
    max_col = ws.max_column or 0
    if max_col <= 0:
        return ''
    rows = []
    if header_row and header_row not in rows:
        rows.append(header_row)
    for r in sorted(set(selected_rows)):
        if r not in rows:
            rows.append(r)
    non_empty_cols = []
    for c in range(1, max_col + 1):
        keep = False
        for r in rows:
            if ws.cell(r, c).value not in (None, ''):
                keep = True
                break
        if keep:
            non_empty_cols.append(c)
    if not non_empty_cols:
        return ''
    merged_map = {}
    skip_cells = set()
    for rng in ws.merged_cells.ranges:
        min_col, min_row, max_col_m, max_row_m = rng.bounds
        if min_row in rows and min_col in non_empty_cols:
            merged_map[(min_row, min_col)] = (max_row_m - min_row + 1, max_col_m - min_col + 1)
        for rr in range(min_row, max_row_m + 1):
            for cc in range(min_col, max_col_m + 1):
                if (rr, cc) != (min_row, min_col):
                    skip_cells.add((rr, cc))
    parts = []
    parts.append("<div style='margin-top:8px; margin-bottom:8px;'>")
    parts.append("<table style='border-collapse:collapse; font-family:Calibri, Microsoft JhengHei, sans-serif; font-size:11pt;'>")
    for idx, r in enumerate(rows):
        row_dim = ws.row_dimensions.get(r)
        row_height = getattr(row_dim, 'height', None)
        tr_style = f"height:{int(row_height)}pt;" if row_height else ""
        parts.append(f"<tr style='{tr_style}'>")
        for c in non_empty_cols:
            if (r, c) in skip_cells:
                continue
            cell = ws.cell(r, c)
            if cell.value is None:
                value = ''
            elif '%' in str(cell.number_format or '') and isinstance(cell.value, (int, float)):
                decimal_match = re.search(r'\.([0#]+)\s*%', str(cell.number_format))
                decimal_places = len(decimal_match.group(1)) if decimal_match else 2
                value = f"{cell.value * 100:.{decimal_places}f}%"
            else:
                value = cell.value
            styles = ["border:1px solid #000", "padding:4px 6px", "vertical-align:middle"]
            fill = getattr(cell.fill, 'fgColor', None)
            fill_rgb = getattr(fill, 'rgb', None) if fill else None
            bg = excel_color_to_css(fill_rgb)
            if bg:
                styles.append(f"background-color:{bg}")
            font = cell.font
            if font:
                if font.bold or idx == 0:
                    styles.append("font-weight:bold")
                if font.italic:
                    styles.append("font-style:italic")
                if font.sz:
                    styles.append(f"font-size:{float(font.sz)}pt")
                fc = excel_color_to_css(getattr(getattr(font, 'color', None), 'rgb', None))
                if fc:
                    styles.append(f"color:{fc}")
            align = cell.alignment
            if align:
                horiz = (align.horizontal or '').lower()
                if horiz in ('center', 'left', 'right', 'justify'):
                    styles.append(f"text-align:{horiz}")
                if align.wrap_text:
                    styles.append("white-space:pre-wrap")
            tag = 'th' if idx == 0 else 'td'
            rowspan, colspan = merged_map.get((r, c), (1, 1))
            span_attr = ''
            if rowspan > 1:
                span_attr += f" rowspan='{rowspan}'"
            if colspan > 1:
                span_attr += f" colspan='{colspan}'"
            style_text = '; '.join(styles)
            parts.append(f"<{tag}{span_attr} style='{style_text}'>{html_escape_keep_breaks(value)}</{tag}>")
        parts.append("</tr>")
    parts.append("</table></div>")
    return ''.join(parts)


def html_table_signature(html_fragment):
    """用於表格去重，避免同一張表被重複貼到信中。"""
    try:
        text = BeautifulSoup(html_fragment or "", "html.parser").get_text(" ", strip=True)
    except Exception:
        text = str(html_fragment or "")
    return re.sub(r"\s+", " ", text).strip().upper()


def group_nearby_rows(row_numbers, max_gap=3):
    """把相近的列分群，避免同一個 sheet 內不同表格黏在一起。"""
    if not row_numbers:
        return []
    normalized = sorted(set(int(r) for r in row_numbers))
    groups = [[normalized[0]]]
    for r in normalized[1:]:
        if r - groups[-1][-1] <= max_gap:
            groups[-1].append(r)
        else:
            groups.append([r])
    return groups


def worksheet_to_html(ws):
    """將 openpyxl 工作表全部內容轉為具有 inline-CSS 的 HTML 表格字串。"""
    max_row = ws.max_row or 0
    max_col = ws.max_column or 0
    if max_row == 0 or max_col == 0:
        return ""

    def is_row_empty(r):
        for c in range(1, max_col + 1):
            v = ws.cell(r, c).value
            if v not in (None, ""):
                return False
        return True

    def is_col_empty(c):
        for r in range(1, max_row + 1):
            v = ws.cell(r, c).value
            if v not in (None, ""):
                return False
        return True

    while max_row > 0 and is_row_empty(max_row):
        max_row -= 1
    while max_col > 0 and is_col_empty(max_col):
        max_col -= 1
    if max_row == 0 or max_col == 0:
        return ""
    merged_map = {}
    skip_cells = set()
    for rng in ws.merged_cells.ranges:
        min_col, min_row, max_col_m, max_row_m = rng.bounds
        merged_map[(min_row, min_col)] = (max_row_m - min_row + 1, max_col_m - min_col + 1)
        for rr in range(min_row, max_row_m + 1):
            for cc in range(min_col, max_col_m + 1):
                if (rr, cc) != (min_row, min_col):
                    skip_cells.add((rr, cc))
    col_width_css = []
    for c in range(1, max_col + 1):
        dim = ws.column_dimensions.get(get_column_letter(c))
        width = getattr(dim, 'width', None)
        px = int((float(width) if width else 12) * 7 + 5)
        col_width_css.append(px)
    parts = []
    parts.append("<div style='margin-top:8px; margin-bottom:8px;'>")
    parts.append("<table style='border-collapse:collapse; font-family:Calibri, Microsoft JhengHei, sans-serif; font-size:11pt;'>")
    parts.append("<colgroup>")
    for px in col_width_css:
        parts.append(f"<col style='width:{px}px;'>")
    parts.append("</colgroup>")
    for r in range(1, max_row + 1):
        row_dim = ws.row_dimensions.get(r)
        row_height = getattr(row_dim, 'height', None)
        tr_style = f"height:{int(row_height)}pt;" if row_height else ""
        parts.append(f"<tr style='{tr_style}'>")
        for c in range(1, max_col + 1):
            if (r, c) in skip_cells:
                continue
            cell = ws.cell(r, c)
            value = cell.value if cell.value is not None else ""
            styles = ["border:1px solid #000", "padding:4px 6px", "vertical-align:middle"]
            fill = getattr(cell.fill, 'fgColor', None)
            fill_rgb = getattr(fill, 'rgb', None) if fill else None
            bg = excel_color_to_css(fill_rgb)
            if bg:
                styles.append(f"background-color:{bg}")
            font = cell.font
            if font:
                if font.bold:
                    styles.append("font-weight:bold")
                if font.italic:
                    styles.append("font-style:italic")
                if font.sz:
                    styles.append(f"font-size:{float(font.sz)}pt")
                fc = excel_color_to_css(getattr(getattr(font, 'color', None), 'rgb', None))
                if fc:
                    styles.append(f"color:{fc}")
            align = cell.alignment
            if align:
                horiz = (align.horizontal or '').lower()
                if horiz in ('center', 'left', 'right', 'justify'):
                    styles.append(f"text-align:{horiz}")
                vert = (align.vertical or '').lower()
                if vert in ('top', 'center', 'bottom'):
                    css_v = 'middle' if vert == 'center' else vert
                    styles.append(f"vertical-align:{css_v}")
                if align.wrap_text:
                    styles.append("white-space:pre-wrap")
            tag = 'th' if r == 1 else 'td'
            rowspan, colspan = merged_map.get((r, c), (1, 1))
            span_attr = ''
            if rowspan > 1:
                span_attr += f" rowspan='{rowspan}'"
            if colspan > 1:
                span_attr += f" colspan='{colspan}'"
            style_text = "; ".join(styles)
            parts.append(f"<{tag}{span_attr} style='{style_text}'>{html.escape(value)}</{tag}>")
        parts.append("</tr>")
    parts.append("</table></div>")
    return ''.join(parts)


def sheet_table_score(ws):
    """評估工作表作為 ELN 資料表的可能性分數，用於挑選最佳 sheet。"""
    non_empty = 0
    header_hits = 0
    max_cols = 0
    max_rows_to_scan = min(ws.max_row or 0, 80)
    max_cols_to_scan = min(ws.max_column or 0, 30)
    header_keywords = [
        "strategy", "counterparty", "product", "event", "event type", "isin",
        "trade", "fixing", "payment", "currency", "denomination", "nominal",
        "coupon", "redemption", "share delivery", "fx", "entitlement"
    ]
    for r in range(1, max_rows_to_scan + 1):
        row_non_empty = 0
        for c in range(1, max_cols_to_scan + 1):
            val = ws.cell(r, c).value
            if val is None:
                continue
            sval = str(val).strip()
            if not sval:
                continue
            non_empty += 1
            row_non_empty += 1
            low = sval.lower()
            if any(k in low for k in header_keywords):
                header_hits += 1
        max_cols = max(max_cols, row_non_empty)
    return non_empty + header_hits * 8 + max_cols * 3, non_empty, header_hits, max_cols


def get_excel_password(issuer, filename=""):
    """依發行機構名稱或檔名從 PASSWORD_MAP 取得對應的 Excel 解密密碼，找不到則回傳 None。"""
    issuer_text = normalize_company_key(issuer)
    filename_upper = str(filename or "").upper()
    if issuer_text == "NOMURA" or "NOMURA" in filename_upper:
        return PASSWORD_MAP.get("NOMURA")
    if issuer_text == "JPM" or any(k in filename_upper for k in ["JPM", "JPMORGAN", "CHASE"]):
        return (
            PASSWORD_MAP.get("JPM")
            or PASSWORD_MAP.get("JPMORGAN")
            or PASSWORD_MAP.get("CHASE")
        )
    return None


def should_force_password_open_for_issuer(issuer, filename=""):
    """判斷指定發行機構的 Excel 是否必須由密碼解密路徑開啟（NOMURA / JPM）。"""
    issuer_text = normalize_company_key(issuer)
    filename_upper = str(filename or "").upper()
    if issuer_text == "NOMURA" or "NOMURA" in filename_upper:
        return True
    if issuer_text == "JPM" or any(k in filename_upper for k in ["JPM", "JPMORGAN", "CHASE"]):
        return True
    return False


def decrypt_excel_to_tempfile(file_path, password):
    """用 msoffcrypto 先把受密碼保護的 Excel 解密到暫存檔，再交給 openpyxl 讀取。"""
    if not password:
        raise ValueError("找不到對應密碼")
    decrypted_buffer = io.BytesIO()
    with open(file_path, "rb") as f:
        office_file = msoffcrypto.OfficeFile(f)
        office_file.load_key(password=password)
        office_file.decrypt(decrypted_buffer)
    fd, temp_output = tempfile.mkstemp(suffix=".xlsx", prefix="eln_decrypted_")
    os.close(fd)
    with open(temp_output, "wb") as out_f:
        out_f.write(decrypted_buffer.getvalue())
    return temp_output


def open_excel_workbook_no_prompt(xlsx_path, issuer="", filename=""):
    """優先直接用 openpyxl 開；若失敗再用 msoffcrypto 先解密後再 load_workbook。回傳 (workbook, temp_path)。"""
    temp_decrypted_path = None
    filename_for_check = filename or os.path.basename(xlsx_path)
    password = get_excel_password(issuer, filename_for_check)
    should_try_password = should_force_password_open_for_issuer(issuer, filename_for_check)
    if not should_try_password:
        try:
            wb = load_workbook(xlsx_path, data_only=False)
            return wb, None
        except Exception as first_error:
            if not password:
                raise first_error
    if not password:
        raise ValueError(f"找不到對應密碼: {filename_for_check}")
    print(f"[INFO] 使用 msoffcrypto 先解密再讀取: {filename_for_check}")
    temp_decrypted_path = decrypt_excel_to_tempfile(xlsx_path, password)
    wb = load_workbook(temp_decrypted_path, data_only=False)
    return wb, temp_decrypted_path


def build_excel_sheet_html(xlsx_path, issuer="", filename="", target_isin="", target_category=None, target_event_type=""):
    """從 Excel 附件中抓出所有符合同一個 ISIN + Event Type/category 的表格區塊。"""
    temp_unlocked_path = None
    try:
        wb, temp_unlocked_path = open_excel_workbook_no_prompt(
            xlsx_path, issuer=issuer, filename=filename,
        )
        matched_sections = []
        seen_signatures = set()
        best_full_ws = None
        best_full_score = -1
        for ws in wb.worksheets:
            score, non_empty, header_hits, max_cols = sheet_table_score(ws)
            if non_empty >= 8 and max_cols >= 4 and score > best_full_score:
                best_full_score = score
                best_full_ws = ws
            max_row = ws.max_row or 0
            max_col = ws.max_column or 0
            if max_row <= 0 or max_col <= 0:
                continue
            if not target_isin:
                continue
            matched_rows = []
            event_explicit_rows = []
            for r in range(1, max_row + 1):
                values = [ws.cell(r, c).value for c in range(1, max_col + 1)]
                joined = row_values_joined(values)
                if not cell_contains_target(joined, target_isin):
                    continue
                matched_rows.append(r)
                if target_event_type or target_category:
                    if row_matches_isin_and_event(
                        joined, target_isin,
                        target_event_type=target_event_type,
                        target_category=target_category,
                        allow_category_fallback=False
                    ):
                        event_explicit_rows.append(r)
            if not matched_rows:
                continue
            row_groups = group_nearby_rows(matched_rows, max_gap=3)
            for group_rows in row_groups:
                header_row = find_header_row_above(ws, group_rows[0], max_col)
                html_fragment = worksheet_selected_rows_to_html(ws, group_rows, header_row=header_row)
                if not html_fragment:
                    continue
                signature = html_table_signature(html_fragment)
                if signature in seen_signatures:
                    continue
                seen_signatures.add(signature)
                explicit_count = sum(1 for r in group_rows if r in event_explicit_rows)
                section_score = score + len(group_rows) * 100 + explicit_count * 200
                section_html = f"<div style='margin-bottom:12px;'>{html_fragment}</div>"
                matched_sections.append((section_score, section_html))
        if matched_sections:
            matched_sections.sort(key=lambda x: x[0], reverse=True)
            return ''.join(html_block for _, html_block in matched_sections)
        if best_full_ws is not None and not target_isin:
            return worksheet_to_html(best_full_ws)
        return ''
    except Exception as e:
        return f"<p>[Excel 表格轉換失敗: {html.escape(str(e))}]</p>"
    finally:
        if temp_unlocked_path and os.path.exists(temp_unlocked_path):
            try:
                os.remove(temp_unlocked_path)
            except Exception:
                pass


def process_jpm_notification_status_pandas(xlsx_path, issuer="", filename="", isin=""):
    """JPMorgan 專用：找到含 'Notification Status' 的表頭，僅保留指定 ISIN 的資料列，輸出為 inline-CSS HTML。"""
    temp_decrypted_path = None
    try:
        try:
            all_sheets = pd.read_excel(xlsx_path, sheet_name=None, header=None, dtype=str, engine='openpyxl', keep_default_na=False)
        except Exception:
            password = get_excel_password(issuer, filename or os.path.basename(xlsx_path))
            if not password:
                print("[INFO] JPM pandas：找不到對應密碼")
                return ""
            temp_decrypted_path = decrypt_excel_to_tempfile(xlsx_path, password)
            all_sheets = pd.read_excel(temp_decrypted_path, sheet_name=None, header=None, dtype=str, engine='openpyxl', keep_default_na=False)
        for sheet_name, df in all_sheets.items():
            df = df.fillna("").astype(str)
            header_idx = None
            for idx, row in df.iterrows():
                if any(str(v).strip() == "Notification Status" for v in row.values):
                    header_idx = idx
                    break
            if header_idx is None:
                continue
            header_vals = [str(v).strip() for v in df.iloc[header_idx].values]
            isin_col_idx = None
            for ci, h in enumerate(header_vals):
                if h.strip().upper() == "ISIN":
                    isin_col_idx = ci
                    break
            isin_upper = str(isin or "").strip().upper()
            data_rows_list = []
            for row_i in range(header_idx + 1, len(df)):
                row_vals = [str(v).strip() for v in df.iloc[row_i].values]
                non_empty = sum(1 for v in row_vals if v and v not in ("nan", "None"))
                if non_empty == 0:
                    break
                if isin_upper and isin_col_idx is not None:
                    cell_isin = row_vals[isin_col_idx] if isin_col_idx < len(row_vals) else ""
                    if cell_isin.upper() != isin_upper:
                        continue
                data_rows_list.append(row_vals)
            if not data_rows_list:
                continue
            num_cols = len(header_vals)
            keep_col_indices = []
            for ci in range(num_cols):
                for row_vals in data_rows_list:
                    val = row_vals[ci] if ci < len(row_vals) else ""
                    if val and val not in ("nan", "None"):
                        keep_col_indices.append(ci)
                        break
            filtered_headers = [header_vals[ci] for ci in keep_col_indices]
            filtered_rows = [
                [row_vals[ci] if ci < len(row_vals) else "" for ci in keep_col_indices]
                for row_vals in data_rows_list
            ]
            parts = [
                "<div style='margin-top:8px; margin-bottom:8px;'>",
                "<table style='border-collapse:collapse; font-family:Calibri,\"Microsoft JhengHei\",sans-serif; font-size:10pt;'>",
                "<tr>",
            ]
            for col in filtered_headers:
                parts.append(
                    f"<th style='border:1px solid #000; padding:4px 8px; text-align:center; "
                    f"font-weight:bold; background-color:#000099; color:#fff;'>"
                    f"{html.escape(str(col))}</th>"
                )
            parts.append("</tr>")
            for row_vals in filtered_rows:
                parts.append("<tr>")
                for val in row_vals:
                    display = val if val not in ("nan", "None") else ""
                    parts.append(
                        f"<td style='border:1px solid #000; padding:4px 8px; text-align:center; color:#000;'>"
                        f"{html.escape(display)}</td>"
                    )
                parts.append("</tr>")
            parts.append("</table></div>")
            result = "".join(parts)
            print(f"[INFO] JPM Notification Status 表格已找到: sheet={sheet_name}, data_rows={len(data_rows_list)}")
            return result
        print("[INFO] JPM pandas：找不到 'Notification Status' 表頭，回退原本邏輯")
        return ""
    except Exception as e:
        print(f"[警告] process_jpm_notification_status_pandas 失敗: {e}")
        return ""
    finally:
        if temp_decrypted_path and os.path.exists(temp_decrypted_path):
            try:
                os.remove(temp_decrypted_path)
            except Exception:
                pass


def build_safe_html_output_path(issuer, isin, filename):
    """將發行機構、ISIN 、檔名組合為安全的輸出 HTML 檔路徑。"""
    issuer_norm = normalize_company_key(issuer)
    safe_issuer = re.sub(r"[^A-Za-z0-9_-]+", "_", issuer_norm or "UNKNOWN")
    safe_isin = re.sub(r"[^A-Za-z0-9_-]+", "_", str(isin or "NOISIN"))
    base_name = os.path.splitext(os.path.basename(filename))[0]
    safe_base_name = re.sub(r"[^A-Za-z0-9_-]+", "_", base_name or "excel")
    html_filename = f"{safe_issuer}_{safe_isin}_{safe_base_name}.html"
    return os.path.join(OUTPUT_DATA_PATH, html_filename)


def save_excel_html_to_output_data(html_content, issuer, isin, filename):
    """將 Excel 表格 HTML 內容寫入 output_data 目錄下的專屬 HTML 檔。"""
    os.makedirs(OUTPUT_DATA_PATH, exist_ok=True)
    output_path = build_safe_html_output_path(issuer, isin, filename)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("<meta charset='utf-8'>\n")
        f.write(html_content)
    print(f"[INFO] Excel HTML 已輸出: {output_path}")
    return output_path


def should_prefer_excel_for_issuer(issuer):
    """判斷指定發行機構是否優先從 Excel 附件擷取內容（JPM、UBS、NOMURA、BBVA）。"""
    issuer_norm = normalize_company_key(issuer)
    return issuer_norm in {"JPM", "UBS", "NOMURA", "BBVA"}


def has_excel_attachment(msg):
    """判斷信件是否含有 Excel 格式（.xlsx / .xlsm / .xls 等）附件。"""
    try:
        count = getattr(msg, "Attachments").Count if getattr(msg, "Attachments", None) is not None else 0
    except Exception:
        count = 0
    for i in range(1, count + 1):
        try:
            filename = str(msg.Attachments.Item(i).FileName or "").lower()
            if filename.endswith((".xlsx", ".xlsm", ".xltx", ".xltm", ".xls")):
                return True
        except Exception:
            pass
    return False


def extract_excel_html_from_mail_attachments(msg, issuer, isin="", target_category=None, target_event_type=""):
    """讀取信件中的 Excel 附件，只保留同一個 ISIN + Event Type 的表格內容。"""
    if not should_prefer_excel_for_issuer(issuer):
        return "", None, []
    if msg.Attachments.Count <= 0:
        return "", None, []
    temp_mail_dir = os.path.join(DRAFT_TEMP_PATH, f"excel_mail_{uuid.uuid4().hex}")
    os.makedirs(temp_mail_dir, exist_ok=True)
    os.makedirs(OUTPUT_DATA_PATH, exist_ok=True)
    saved_excel_files = []
    collected_fragments = []
    seen_signatures = set()
    issuer_norm = normalize_company_key(issuer)
    for i in range(1, msg.Attachments.Count + 1):
        try:
            att = msg.Attachments.Item(i)
            filename = str(att.FileName or "")
            lower_name = filename.lower()
            if not lower_name.endswith((".xlsx", ".xlsm", ".xltx", ".xltm", ".xls")):
                continue
            excel_path = os.path.join(temp_mail_dir, filename)
            att.SaveAsFile(excel_path)
            saved_excel_files.append(excel_path)
            if issuer_norm == "JPM":
                excel_html = process_jpm_notification_status_pandas(
                    excel_path, issuer=issuer, filename=filename, isin=isin
                )
                if not excel_html.strip():
                    print(f"[INFO] JPM Notification Status 未找到，回退 build_excel_sheet_html: {filename}")
                    excel_html = build_excel_sheet_html(
                        excel_path, issuer=issuer, filename=filename,
                        target_isin=isin, target_category=target_category, target_event_type=target_event_type,
                    )
            else:
                excel_html = build_excel_sheet_html(
                    excel_path, issuer=issuer, filename=filename,
                    target_isin=isin, target_category=target_category, target_event_type=target_event_type,
                )
            if not excel_html.strip():
                continue
            signature = html_table_signature(excel_html)
            if signature in seen_signatures:
                continue
            seen_signatures.add(signature)
            save_excel_html_to_output_data(excel_html, issuer, isin, filename)
            collected_fragments.append(excel_html)
        except Exception as e:
            print(f"[警告] Excel 附件處理失敗: {e}")
    merged_html = "".join(collected_fragments)
    return merged_html, temp_mail_dir, saved_excel_files


def get_pdf_matching_pages(pdf_path, target_isin, target_category=None, target_event_type=""):
    """掃描 PDF 各頁文字，回傳包含目標 ISIN（且符合事件分類）的頁碼列表。"""
    matched_pages = []
    doc = fitz.open(pdf_path)
    try:
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            text = page.get_text('text') or ''
            if not cell_contains_target(text, target_isin):
                continue
            if target_event_type or target_category:
                if not row_matches_isin_and_event(text, target_isin, target_event_type=target_event_type, target_category=target_category):
                    continue
            matched_pages.append(page_index)
    finally:
        doc.close()
    return matched_pages


def save_pdf_selected_pages_as_images(pdf_path, output_dir, page_indexes, zoom=1.8, prefix='page'):
    """將 PDF 中指定頁碼渲染成 PNG 圖片儲存至輸出目錄，回傳圖片路徑列表。"""
    image_paths = []
    if not page_indexes:
        return image_paths
    os.makedirs(output_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    try:
        for page_index in page_indexes:
            page = doc.load_page(page_index)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            out_path = os.path.join(output_dir, f"{prefix}_{page_index + 1}.png")
            pix.save(out_path)
            image_paths.append(out_path)
    finally:
        doc.close()
    return image_paths


def save_pdf_pages_as_images(pdf_path, output_dir, zoom=1.8, prefix="page"):
    """將 PDF 全部頁面渲染成 PNG 圖片（BNP 專用），回傳圖片路徑列表。"""
    image_paths = []
    os.makedirs(output_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    try:
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            out_path = os.path.join(output_dir, f"{prefix}_{page_index + 1}.png")
            pix.save(out_path)
            image_paths.append(out_path)
    finally:
        doc.close()
    return image_paths


def attach_inline_image(mail, image_path):
    """將圖片以 inline 方式嵌入 Outlook 信件，設定 CID 屬性後回傳 CID 供 HTML <img> 引用。"""
    attachment = mail.Attachments.Add(image_path, 1, 0)
    cid = str(uuid.uuid4())
    mail.Save()
    pa = attachment.PropertyAccessor
    pa.SetProperty("http://schemas.microsoft.com/mapi/proptag/0x3712001F", cid)
    pa.SetProperty("http://schemas.microsoft.com/mapi/proptag/0x3713001F", cid)
    return cid


def extract_hsbc_conversion_text_block(text, target_isin):
    """HSBC conversion：從 1) 開始，抓到表格前的文字內容，排除表格欄位名稱與資料列。"""
    if not text:
        return ""
    raw_lines = str(text).splitlines()
    lines = [line.rstrip() for line in raw_lines]
    numbered_re = re.compile(r'^\s*1\)')
    start_idx = None
    for i, line in enumerate(lines):
        if cell_contains_target(line, target_isin):
            for j in range(max(0, i - 10), i + 1):
                if numbered_re.match(lines[j].strip()):
                    start_idx = j
                    break
            if start_idx is not None:
                break
    if start_idx is None:
        for i, line in enumerate(lines):
            if numbered_re.match(line.strip()):
                start_idx = i
                break
    if start_idx is None:
        return ""
    collected = []
    footer_keywords = ('regards', 'best regards', 'thanks', 'thank you', 'sent from')
    table_header_tokens = {'isin', 'fx', 'strike', 'closing'}
    for idx in range(start_idx, len(lines)):
        line = lines[idx]
        stripped = line.strip()
        low = stripped.lower()
        if stripped and any(low.startswith(k) for k in footer_keywords):
            break
        if not stripped:
            collected.append(line)
            continue
        if low in table_header_tokens:
            break
        upcoming = [
            lines[k].strip().lower()
            for k in range(idx, min(idx + 4, len(lines)))
            if lines[k].strip()
        ]
        if upcoming[:4] == ['isin', 'fx', 'strike', 'closing']:
            break
        if cell_contains_target(stripped, target_isin):
            next_nonempty = []
            for k in range(idx + 1, min(idx + 6, len(lines))):
                v = lines[k].strip()
                if v:
                    next_nonempty.append(v)
            if next_nonempty and any(re.fullmatch(r'[\d.,]+', x) for x in next_nonempty):
                break
        collected.append(line)
    while collected and not collected[0].strip():
        collected.pop(0)
    while collected and not collected[-1].strip():
        collected.pop()
    if not collected:
        return ""
    html_parts = []
    for line in collected:
        stripped = line.strip()
        if not stripped:
            html_parts.append("<div style='height:10px;'></div>")
        else:
            html_parts.append(
                f"<div style='margin:0 0 8px 0; line-height:1.35; color:#000;'>"
                f"{html.escape(stripped)}"
                f"</div>"
            )
    return "".join(html_parts)


def extract_hsbc_conversion_html_fragment(html_body, target_isin):
    """HSBC conversion：抓取文字（從 1) 到表格前）與含目標 ISIN 的表格，合併輸出。"""
    if not html_body:
        return ""
    try:
        soup = BeautifulSoup(html_body, "html.parser")
    except Exception:
        return ""
    text_part = extract_hsbc_conversion_text_block(
        soup.get_text("\n", strip=False), target_isin
    )
    chosen_table_html = ""
    for table in soup.find_all("table"):
        table_text = table.get_text(" ", strip=True)
        if not cell_contains_target(table_text, target_isin):
            continue
        table_soup = BeautifulSoup(str(table), "html.parser")
        for a in table_soup.find_all("a"):
            a.replace_with(a.get_text(" ", strip=True))
        for tag in table_soup.find_all(["table", "tr", "td", "th"]):
            if tag.name == "table":
                tag.attrs = {"style": "border-collapse:collapse; font-family:Calibri, Microsoft JhengHei, sans-serif; font-size:11pt; color:#000;"}
            elif tag.name == "th":
                tag.attrs = {"style": "border:1px solid #000; padding:4px 8px; text-align:center; font-weight:bold; background-color:#fff200; color:#000;"}
            elif tag.name == "td":
                tag.attrs = {"style": "border:1px solid #000; padding:4px 8px; text-align:center; color:#000;"}
            elif tag.name == "tr":
                tag.attrs = {}
        chosen_table_html = str(table_soup)
        break
    if text_part and chosen_table_html:
        return (
            "<div style='margin-top:8px; margin-bottom:8px;'>"
            f"{text_part}"
            "<div style='height:10px;'></div>"
            f"{chosen_table_html}"
            "</div>"
        )
    if chosen_table_html:
        return f"<div style='margin-top:8px; margin-bottom:8px;'>{chosen_table_html}</div>"
    return text_part


def extract_sg_table_only_html_fragment(html_body, target_isin, target_category=None, target_event_type=""):
    """SG / SGMARKETS：只抓最內層、真正有資料的表格，排除外層排版 table、footer、免責聲明。"""
    if not html_body:
        return ""
    try:
        soup = BeautifulSoup(html_body, "html.parser")
    except Exception:
        return ""

    def is_footer_or_layout_table(table_text):
        low = re.sub(r'\s+', ' ', str(table_text or '').strip()).lower()
        if not low:
            return True
        # If the table contains core data-table keywords it IS a data table, never a footer
        data_table_kws = ['event type', 'payment date', 'valuation date', 'settlement type']
        if sum(1 for k in data_table_kws if k in low) >= 2:
            return False
        footer_keywords = [
            'best regards', 'contact us', 'view my products', 'indicative document',
            'no offer to contract', 'confidentiality', 'market information',
            'société générale', 'societe generale', 'sg markets life cycle team', 'dear 300operations'
        ]
        return any(k in low for k in footer_keywords)

    def row_cell_texts(tr):
        cells = tr.find_all(["th", "td"], recursive=False)
        if not cells:
            cells = tr.find_all(["th", "td"])
        return [c.get_text(" ", strip=True) for c in cells]

    def normalize_header_cells(cells):
        return [re.sub(r'\s+', ' ', c).strip().lower() for c in cells if str(c).strip()]

    def looks_like_data_table(table):
        rows = table.find_all('tr')
        if not rows:
            return False
        max_cols = 0
        header_hits = 0
        for tr in rows[:6]:
            cells = row_cell_texts(tr)
            max_cols = max(max_cols, len(cells))
            norm = normalize_header_cells(cells)
            joined = ' | '.join(norm)
            for kw in ['event type', 'isin', 'settlement type', 'valuation date', 'payment date',
                       'maturity date', 'coupon amount', 'delivered underlying',
                       'underlying isin', 'underlying ticker', 'number of shares per note']:
                if kw in joined:
                    header_hits += 1
        return max_cols >= 4 and header_hits >= 1

    scored_tables = []
    seen_signatures = set()
    all_tables = soup.find_all('table')
    leaf_tables = [t for t in all_tables if not t.find('table')]

    def get_direct_rows(tbl):
        """Return only <tr> rows that directly belong to tbl, not to any nested sub-table."""
        return [tr for tr in tbl.find_all('tr')
                if next((p for p in tr.parents if p.name == 'table'), None) is tbl]

    # First pass: leaf tables (no nested sub-tables).  Second pass (fallback): non-leaf tables
    # using direct-row iteration so inner sub-table rows are excluded.
    passes = [(leaf_tables, False)]
    non_leaf_tables = [t for t in all_tables if t.find('table')]
    if non_leaf_tables:
        passes.append((non_leaf_tables, True))

    for table_list, use_direct_rows in passes:
        if scored_tables:
            break
        for table in table_list:
            table_text = table.get_text(' ', strip=True)
            if not cell_contains_target(table_text, target_isin):
                continue
            if is_footer_or_layout_table(table_text):
                continue
            if not looks_like_data_table(table):
                continue
            rows = get_direct_rows(table) if use_direct_rows else table.find_all('tr')
            header_cells = []
            matched_rows = []
            explicit_event_rows = 0
            _hdr_kws = ['event type', 'isin', 'settlement type', 'valuation date', 'payment date',
                        'maturity date', 'coupon amount', 'delivered underlying',
                        'underlying isin', 'underlying ticker', 'number of shares per note']
            # Non-leaf (second pass): the outer table's direct rows may include a single-cell
            # wrapper <tr> that wraps an inner leaf table containing the real column headers.
            # Pre-scan all nested sub-tables to extract the column header row before the main
            # loop, so the wrapper row is never mis-captured as header_cells.
            if use_direct_rows:
                for sub_tbl in table.find_all('table'):
                    for n_tr in sub_tbl.find_all('tr')[:6]:
                        n_cells = row_cell_texts(n_tr)
                        if len(n_cells) >= 2:
                            n_norm = normalize_header_cells(n_cells)
                            if (n_tr.find('th', recursive=False) or
                                    any(k in ' | '.join(n_norm) for k in _hdr_kws)):
                                header_cells = n_cells
                                break
                    if header_cells:
                        break
            for idx, tr in enumerate(rows):
                cells = row_cell_texts(tr)
                if not cells:
                    continue
                row_text = ' '.join(cells)
                norm_cells = normalize_header_cells(cells)
                # len(cells) >= 2 prevents a single-cell wrapper <tr> from being captured
                # as the header; recursive=False prevents nested <th> from firing here.
                if idx <= 6 and not header_cells and len(cells) >= 2 and (tr.find('th', recursive=False) or any(k in ' | '.join(norm_cells) for k in _hdr_kws)):
                    header_cells = cells
                    continue
                if not cell_contains_target(row_text, target_isin):
                    continue
                if target_event_type or target_category:
                    if row_matches_isin_and_event(row_text, target_isin, target_event_type=target_event_type, target_category=target_category, allow_category_fallback=False):
                        explicit_event_rows += 1
                        matched_rows.append(cells)
                    elif not explicit_event_rows:
                        matched_rows.append(cells)
                else:
                    matched_rows.append(cells)
            if not matched_rows:
                continue
            # Skip tables where any data cell is unreasonably long (indicates email body text, not structured data)
            if any(len(str(cell)) > 300 for row in matched_rows for cell in row):
                continue
            parts = ["<table style='border-collapse:collapse; font-family:Calibri, Microsoft JhengHei, sans-serif; font-size:11pt; color:#000; margin-top:0; margin-bottom:0;'>"]
            if header_cells:
                parts.append('<tr>')
                for cell in header_cells:
                    parts.append(f"<th style='border:1px solid #000; padding:4px 6px; text-align:center; font-weight:bold; background-color:#000099; color:#fff;'>{html.escape(str(cell))}</th>")
                parts.append('</tr>')
            for row in matched_rows:
                parts.append('<tr>')
                for cell in row:
                    parts.append(f"<td style='border:1px solid #000; padding:4px 6px; text-align:center; color:#000;'>{html.escape(str(cell))}</td>")
                parts.append('</tr>')
            parts.append('</table>')
            table_html = ''.join(parts)
            signature = html_table_signature(table_html)
            if signature in seen_signatures:
                continue
            seen_signatures.add(signature)
            score = explicit_event_rows * 300 + len(matched_rows) * 100
            if target_event_type or target_category:
                if text_matches_target_event_type(table_text, target_event_type=target_event_type, target_category=target_category):
                    score += 80
            if header_cells:
                score += 30
            scored_tables.append((score, table_html))
    if not scored_tables:
        return ""
    scored_tables.sort(key=lambda x: x[0], reverse=True)
    top_tables = scored_tables[:3]
    return ''.join([f"<div style='margin-bottom:12px;'>{tbl}</div>" for _, tbl in top_tables])


def extract_barclays_html_fragment(html_body, target_isin):
    """Barclays 專用：從信件 HTML body 找到含目標 ISIN 的表格，套用乾淨 inline-CSS 輸出。"""
    if not html_body:
        return ""
    try:
        soup = BeautifulSoup(html_body, "html.parser")
    except Exception:
        return ""
    isin_upper = str(target_isin or "").strip().upper()
    for table in soup.find_all("table"):
        table_text = table.get_text(" ", strip=True)
        if not cell_contains_target(table_text, isin_upper):
            continue
        rows = table.find_all("tr")
        if not rows:
            continue
        header_rows_cells = []
        data_rows_cells = []
        for tr in rows:
            cells = tr.find_all(["th", "td"])
            cell_texts = [c.get_text(" ", strip=True) for c in cells]
            if tr.find("th") or any(t.strip().upper() == "ISIN" for t in cell_texts):
                header_rows_cells.append(cells)
            else:
                row_text = " ".join(cell_texts)
                if cell_contains_target(row_text, isin_upper):
                    data_rows_cells.append(cells)
        if not data_rows_cells:
            continue
        title_text = "Fixing Events Notification"
        parts = [
            "<div style='margin-top:8px; margin-bottom:8px;'>",
            f"<div style='font-weight:bold; font-size:11pt; font-family:Calibri,\"Microsoft JhengHei\",sans-serif; margin-bottom:6px; color:#000;'>{html.escape(title_text)}</div>",
            "<table style='border-collapse:collapse; font-family:Calibri,\"Microsoft JhengHei\",sans-serif; font-size:10pt;'>",
        ]
        for cells in header_rows_cells:
            parts.append("<tr>")
            for cell in cells:
                cs = cell.get("colspan", "")
                rs = cell.get("rowspan", "")
                span = (f" colspan='{cs}'" if cs and cs != "1" else "") + (f" rowspan='{rs}'" if rs and rs != "1" else "")
                parts.append(f"<th{span} style='border:1px solid #ccc; padding:4px 8px; text-align:center; font-weight:bold; background-color:#1f3864; color:#fff;'>{html.escape(cell.get_text(' ', strip=True))}</th>")
            parts.append("</tr>")
        for cells in data_rows_cells:
            parts.append("<tr>")
            for cell in cells:
                cs = cell.get("colspan", "")
                rs = cell.get("rowspan", "")
                span = (f" colspan='{cs}'" if cs and cs != "1" else "") + (f" rowspan='{rs}'" if rs and rs != "1" else "")
                parts.append(f"<td{span} style='border:1px solid #ccc; padding:4px 8px; text-align:center; color:#000;'>{html.escape(cell.get_text(' ', strip=True))}</td>")
            parts.append("</tr>")
        parts.append("</table></div>")
        return "".join(parts)
    print("[INFO] Barclays：找不到含 ISIN 的表格")
    return ""


def extract_barclays_ticker_from_html(html_body, target_isin=""):
    if not html_body:
        return ""
    try:
        soup = BeautifulSoup(html_body, "html.parser")
    except Exception:
        return ""
    target_isin_upper = str(target_isin or "").strip().upper()
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        header_row = None
        note_isin_col_index = -1
        stock_col_index = -1
        for tr in rows:
            cells = tr.find_all(["th", "td"])
            if not cells:
                continue
            stock_idx = find_header_cell_index_by_keywords(cells, [["stock", "redemption", "sedol", "isin"]])
            if stock_idx == -1:
                continue
            for ci, cell in enumerate(cells):
                if normalize_html_column_name(cell.get_text(" ", strip=True)) == "isin":
                    note_isin_col_index = ci
                    stock_col_index = stock_idx
                    header_row = tr
                    break
            if header_row is not None:
                break
        if header_row is None or stock_col_index == -1:
            continue
        passed_header = False
        for tr in rows:
            if tr == header_row:
                passed_header = True
                continue
            if not passed_header:
                continue
            cells = tr.find_all(["th", "td"])
            if len(cells) <= stock_col_index:
                continue
            row_text_upper = tr.get_text(" ", strip=True).upper()
            if target_isin_upper:
                row_note_isin = ""
                if note_isin_col_index != -1 and len(cells) > note_isin_col_index:
                    row_note_isin = cells[note_isin_col_index].get_text(" ", strip=True).upper()
                if target_isin_upper not in row_note_isin and target_isin_upper not in row_text_upper:
                    continue
            stock_text = cells[stock_col_index].get_text(" ", strip=True)
            ticker = extract_barclays_stock_ticker_from_text(stock_text)
            if ticker:
                return ticker
    return ""


def lookup_barclays_ticker_from_outlook(start_date, end_date, target_isin):
    target_isin_upper = str(target_isin or "").strip().upper()
    if not target_isin_upper:
        return ""
    try:
        outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
        target_folder = get_target_outlook_folder(
            outlook_namespace=outlook,
            account_name=OUTLOOK_ACCOUNT,
            target_folder_name=TARGET_FOLDER_NAME
        )
        messages = get_outlook_messages(target_folder, start_date, end_date)
        for msg in messages:
            try:
                if msg.Class != 43:
                    continue
                subject = str(getattr(msg, "Subject", "") or "")
                body = str(getattr(msg, "Body", "") or "")
                html_body = str(getattr(msg, "HTMLBody", "") or "")
                blob_upper = "\n".join([subject, body, html_body]).upper()
                if target_isin_upper not in blob_upper:
                    continue
                sender_name = getattr(msg, "SenderName", "")
                try:
                    sender_email = msg.SenderEmailAddress
                except Exception:
                    sender_email = ""
                extracted_company = extract_company_domain(sender_email, sender_name)
                if "BARCLAYS" not in extracted_company and "BARCLAYS" not in subject.upper():
                    continue
                ticker = extract_barclays_ticker_from_html(html_body, target_isin_upper)
                if ticker:
                    return ticker
            except Exception:
                continue
    except Exception as e:
        print(f"[WARN] Barclays Outlook ticker lookup failed ({target_isin_upper}): {e}")
    return ""


def extract_natixis_html_fragment(html_body, target_isin, target_event_type=""):
    """Natixis 專用：從信件 HTML body 找到含目標 ISIN 的表格，套用乾淨 inline-CSS 輸出。"""
    if not html_body:
        return ""
    try:
        soup = BeautifulSoup(html_body, "html.parser")
    except Exception:
        return ""
    isin_upper = str(target_isin or "").strip().upper()
    target_event_norm = str(target_event_type or "").strip().upper()
    for table in soup.find_all("table"):
        table_text = table.get_text(" ", strip=True)
        if not cell_contains_target(table_text, isin_upper):
            continue
        rows = table.find_all("tr")
        if not rows:
            continue
        header_rows_cells = []
        data_rows_cells = []
        redemption_col_idx = -1
        for tr in rows:
            cells = tr.find_all(["th", "td"])
            cell_texts = [c.get_text(" ", strip=True) for c in cells]
            cell_upper = [t.strip().upper() for t in cell_texts]
            is_header = tr.find("th") or any(t == "ISIN" for t in cell_upper)
            if is_header:
                if redemption_col_idx == -1:
                    for ci, t in enumerate(cell_upper):
                        if t == "REDEMPTION":
                            redemption_col_idx = ci
                            break
                header_rows_cells.append(cells)
            else:
                row_text = " ".join(cell_texts)
                if not cell_contains_target(row_text, isin_upper):
                    continue
                if target_event_norm and redemption_col_idx != -1 and redemption_col_idx < len(cells):
                    redemption_val = cells[redemption_col_idx].get_text(" ", strip=True).strip().upper()
                    if redemption_val and target_event_norm not in redemption_val and redemption_val not in target_event_norm:
                        continue
                data_rows_cells.append(cells)
        if not data_rows_cells:
            continue
        parts = [
            "<div style='margin-top:8px; margin-bottom:8px;'>",
            "<table style='border-collapse:collapse; font-family:Calibri,\"Microsoft JhengHei\",sans-serif; font-size:10pt;'>",
        ]
        for cells in header_rows_cells:
            parts.append("<tr>")
            for cell in cells:
                cs = cell.get("colspan", "")
                rs = cell.get("rowspan", "")
                span = (f" colspan='{cs}'" if cs and cs != "1" else "") + (f" rowspan='{rs}'" if rs and rs != "1" else "")
                parts.append(f"<th{span} style='border:1px solid #000; padding:4px 8px; text-align:center; font-weight:bold; background-color:#6B0080; color:#fff;'>{html.escape(cell.get_text(' ', strip=True))}</th>")
            parts.append("</tr>")
        for cells in data_rows_cells:
            parts.append("<tr>")
            for cell in cells:
                cs = cell.get("colspan", "")
                rs = cell.get("rowspan", "")
                span = (f" colspan='{cs}'" if cs and cs != "1" else "") + (f" rowspan='{rs}'" if rs and rs != "1" else "")
                parts.append(f"<td{span} style='border:1px solid #000; padding:4px 8px; text-align:center; color:#000;'>{html.escape(cell.get_text(' ', strip=True))}</td>")
            parts.append("</tr>")
        parts.append("</table></div>")
        result = "".join(parts)
        print(f"[INFO] Natixis 表格已找到: data_rows={len(data_rows_cells)}")
        return result
    print("[INFO] Natixis：找不到含 ISIN 的表格")
    return ""


def extract_matching_text_lines(text, target_isin, target_category=None, target_event_type=""):
    """從純文字信件中擷取包含目標 ISIN 的相關行，輸出為 <p> HTML 段落。"""
    if not text:
        return ''
    lines = [line.strip() for line in str(text).splitlines() if line.strip()]
    matched = []
    for i, line in enumerate(lines):
        if not cell_contains_target(line, target_isin):
            continue
        start = max(0, i - 1)
        end = min(len(lines), i + 2)
        snippet = lines[start:end]
        joined = ' '.join(snippet)
        if target_event_type or target_category:
            if not row_matches_isin_and_event(joined, target_isin, target_event_type=target_event_type, target_category=target_category):
                continue
        matched.extend(snippet)
    seen = []
    for line in matched:
        if line not in seen:
            seen.append(line)
    return ''.join([f"<p>{html.escape(line)}</p>" for line in seen])


def extract_matching_html_fragment(html_body, target_isin, target_category=None, target_event_type=""):
    """從信件 HTML 內文中抓出所有符合同一個 ISIN + Event Type/category 的表格。"""
    if not html_body:
        return ''
    try:
        soup = BeautifulSoup(html_body, 'html.parser')
    except Exception:
        return ''
    scored_tables = []
    seen_signatures = set()
    for table in soup.find_all('table'):
        rows = table.find_all('tr')
        if not rows:
            continue
        header_html = ''
        isin_rows = []
        event_matched_rows = []
        table_text = table.get_text(' ', strip=True)
        table_has_target_event = (
            text_matches_target_event_type(table_text, target_event_type=target_event_type, target_category=target_category)
            if (target_event_type or target_category) else False
        )
        
        # 智能檢測表頭行：尋找包含關鍵欄位名稱的行
        header_keywords = ['isin', 'event', 'payment', 'fixing', 'denomination', 'currency', 'redemption', 'coupon', 'shares']
        header_row_idx = -1
        
        for idx, tr in enumerate(rows):
            cells = tr.find_all(['th', 'td'])
            row_text = ' '.join(c.get_text(' ', strip=True) for c in cells).lower()
            
            # 跳過包含說明文字的行
            if any(phrase in row_text for phrase in ['hi all', 'please find', 'thank you', 'the indicative', 'final version']):
                continue
            
            # 檢查是否為表頭（包含多個關鍵欄位名稱）
            if header_row_idx == -1:
                keyword_count = sum(1 for kw in header_keywords if kw in row_text)
                if keyword_count >= 3:  # 至少包含3個關鍵欄位名稱
                    header_html = str(tr)
                    header_row_idx = idx
                    continue
        
        for idx, tr in enumerate(rows):
            # 跳過已識別的表頭行
            if idx == header_row_idx:
                continue
                
            cells = tr.find_all(['th', 'td'])
            row_text = ' '.join(c.get_text(' ', strip=True) for c in cells)
            
            # 跳過說明文字行
            if any(phrase in row_text.lower() for phrase in ['hi all', 'please find', 'thank you', 'the indicative', 'final version']):
                continue
            if not cell_contains_target(row_text, target_isin):
                continue
            isin_rows.append(str(tr))
            if target_event_type or target_category:
                if row_matches_isin_and_event(row_text, target_isin, target_event_type=target_event_type, target_category=target_category, allow_category_fallback=False):
                    event_matched_rows.append(str(tr))
                elif not event_matched_rows:
                    event_matched_rows.append(str(tr))
            else:
                event_matched_rows.append(str(tr))
        chosen_rows = event_matched_rows if event_matched_rows else isin_rows
        if not chosen_rows:
            continue
        score = len(event_matched_rows) * 300 + len(isin_rows) * 100 + (50 if table_has_target_event else 0)
        if header_html and (target_event_type or target_category):
            if text_matches_target_event_type(header_html, target_event_type=target_event_type, target_category=target_category):
                score += 80
        table_html = "<table style='border-collapse: collapse;'>"
        if header_html:
            table_html += header_html
        table_html += ''.join(chosen_rows)
        table_html += '</table>'
        signature = html_table_signature(table_html)
        if signature in seen_signatures:
            continue
        seen_signatures.add(signature)
        scored_tables.append((score, table_html))
    if not scored_tables:
        return ''
    scored_tables.sort(key=lambda x: x[0], reverse=True)
    return ''.join([f"<div style='margin-bottom:12px;'>{tbl}</div>" for _, tbl in scored_tables])


def get_mail_content_blocks(start_date, end_date, issuer, isin, target_category=None, target_event_type=""):
    """依 issuer + ISIN + Event Type 回 Outlook 找對應原始信件，擷取符合條件的內容區塊。"""
    content_blocks = []
    try:
        outlook_ns = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
        target_folder = get_target_outlook_folder(
            outlook_namespace=outlook_ns,
            account_name=OUTLOOK_ACCOUNT,
            target_folder_name=TARGET_FOLDER_NAME
        )
        messages = get_outlook_messages(target_folder, start_date, end_date)
    except Exception as e:
        print(f"[警告] 讀取原始信件失敗: {e}")
        return content_blocks
    issuer_norm = normalize_company_key(issuer)
    isin_upper = str(isin or "").upper().strip()
    normalized_target_category = target_category or normalize_event_type_to_category(target_event_type)
    seen_keys = set()
    for msg in messages:
        try:
            if msg.Class != 43:
                continue
            mail_date = msg.ReceivedTime.date()
            if mail_date < start_date or mail_date > end_date:
                continue
            sender_name = getattr(msg, "SenderName", "")
            try:
                sender_email = msg.SenderEmailAddress
            except Exception:
                sender_email = ""
            extracted_company = extract_company_domain(sender_email, sender_name)
            subject = str(getattr(msg, "Subject", "") or "")
            body = str(getattr(msg, "Body", "") or "")
            html_body = str(getattr(msg, "HTMLBody", "") or "")
            if not is_company_match(issuer_norm, extracted_company, sender_name, subject):
                continue
            # SG 只接受主旨為 "Events on your Structured Products" 的信件
            if issuer_norm == "SG" and "Events on your Structured Products" not in subject:
                print(f"[SG 跳過] 步驟4主旨不符，略過此信: {subject}")
                continue
            blob = "\n".join([subject, body, html_body]).upper()
            attachment_names = []
            if msg.Attachments.Count > 0:
                for i in range(1, msg.Attachments.Count + 1):
                    try:
                        attachment_names.append(str(msg.Attachments.Item(i).FileName))
                    except Exception:
                        pass
            issuer_prefers_excel = should_prefer_excel_for_issuer(issuer_norm)
            excel_attachment_exists = has_excel_attachment(msg)
            isin_found = True
            if isin_upper:
                isin_found = (isin_upper in blob) or any(isin_upper in x.upper() for x in attachment_names)
            if isin_upper and (not isin_found):
                if issuer_prefers_excel and excel_attachment_exists:
                    print(f"[INFO] {issuer_norm} 偵測到 Excel 附件，直接抓取: {subject}")
                else:
                    continue
            joined_message_text = "\n".join([subject, body, BeautifulSoup(html_body, "html.parser").get_text(" ", strip=True)])
            if normalized_target_category:
                if not content_text_matches_category(joined_message_text, normalized_target_category):
                    html_has_match = False
                    try:
                        if html_body:
                            fragment = extract_matching_html_fragment(html_body, isin_upper, normalized_target_category, target_event_type=target_event_type)
                            if fragment:
                                html_has_match = True
                    except Exception:
                        html_has_match = False
                    if not html_has_match and not issuer_prefers_excel and issuer_norm not in ("BNP", "SG"):
                        continue
            unique_key = f"{msg.EntryID}|{isin_upper}|{normalized_target_category}"
            if unique_key in seen_keys:
                continue
            seen_keys.add(unique_key)
            block = {
                "title": subject,
                "company": extracted_company,
                "received": msg.ReceivedTime.strftime("%Y-%m-%d %H:%M"),
                "html_fragment": "",
                "pdf_images": [],
                "excel_files": [],
                "temp_dir": None,
            }
            pdf_found = False
            excel_found = False
            if should_prefer_excel_for_issuer(issuer_norm):
                excel_html, excel_temp_dir, excel_files = extract_excel_html_from_mail_attachments(
                    msg, issuer_norm, isin_upper,
                    target_category=normalized_target_category,
                    target_event_type=target_event_type,
                )
                if excel_temp_dir:
                    block["temp_dir"] = excel_temp_dir
                if excel_files:
                    block["excel_files"] = excel_files
                if excel_html:
                    excel_found = True
                    block["html_fragment"] = excel_html
            if (not excel_found) and msg.Attachments.Count > 0:
                temp_mail_dir = block["temp_dir"] or os.path.join(DRAFT_TEMP_PATH, f"mail_{uuid.uuid4().hex}")
                os.makedirs(temp_mail_dir, exist_ok=True)
                block["temp_dir"] = temp_mail_dir
                output_mail_dir = os.path.join(OUTPUT_DATA_PATH, f"mail_{uuid.uuid4().hex}")
                os.makedirs(output_mail_dir, exist_ok=True)
                for i in range(1, msg.Attachments.Count + 1):
                    try:
                        att = msg.Attachments.Item(i)
                        filename = str(att.FileName)
                        lower_name = filename.lower()
                        if lower_name.endswith(".pdf"):
                            pdf_path = os.path.join(temp_mail_dir, filename)
                            att.SaveAsFile(pdf_path)
                            safe_prefix = re.sub(r"[^A-Za-z0-9_-]+", "_", os.path.splitext(filename)[0]) or "page"
                            if issuer_norm == "BNP":
                                pdf_found = True
                                image_paths = save_pdf_pages_as_images(pdf_path, output_mail_dir, prefix=safe_prefix)
                            else:
                                page_indexes = get_pdf_matching_pages(pdf_path, isin_upper, normalized_target_category, target_event_type=target_event_type)
                                if not page_indexes:
                                    continue
                                pdf_found = True
                                image_paths = save_pdf_selected_pages_as_images(pdf_path, output_mail_dir, page_indexes, prefix=safe_prefix)
                            block["pdf_images"].extend(image_paths)
                    except Exception as e:
                        print(f"[警告] PDF 附件處理失敗: {e}")
            if not excel_found and not pdf_found:
                html_fragment = ""
                if issuer_norm == "HSBC" and normalized_target_category == "conversion":
                    html_fragment = extract_hsbc_conversion_html_fragment(html_body, isin_upper)
                    if not html_fragment:
                        html_fragment = extract_hsbc_conversion_text_block(body, isin_upper)
                elif issuer_norm == "SG":
                    html_fragment = extract_sg_table_only_html_fragment(
                        html_body, isin_upper, normalized_target_category, target_event_type=target_event_type,
                    )
                elif issuer_norm == "BARCLAYS":
                    html_fragment = extract_barclays_html_fragment(html_body, isin_upper)
                elif issuer_norm == "NATIXIS":
                    html_fragment = extract_natixis_html_fragment(html_body, isin_upper, target_event_type=target_event_type)
                if not html_fragment and issuer_norm not in ("NATIXIS", "SG"):
                    html_fragment = extract_matching_html_fragment(
                        html_body, isin_upper, normalized_target_category, target_event_type=target_event_type,
                    )
                if not html_fragment and issuer_norm not in ("SG", "NATIXIS"):
                    html_fragment = extract_matching_text_lines(
                        body, isin_upper, normalized_target_category, target_event_type=target_event_type,
                    )
                if html_fragment:
                    block["html_fragment"] = html_fragment
            if block["html_fragment"] or block["pdf_images"] or block["excel_files"]:
                content_blocks.append(block)
            else:
                if block.get("temp_dir") and os.path.exists(block["temp_dir"]):
                    shutil.rmtree(block["temp_dir"], ignore_errors=True)
        except Exception as e:
            print(f"[警告] 比對原始信件時略過一封郵件: {e}")
            continue
    return content_blocks


def build_source_content_html(mail, start_date, end_date, issuer, isin, target_category=None, target_event_type=""):
    """組合 Outlook 原始信件內容（表格/PDF 圖片）為 HTML，嵌入草稿信中。"""
    blocks = get_mail_content_blocks(
        start_date, end_date, issuer, isin,
        target_category=target_category,
        target_event_type=target_event_type,
    )
    if not blocks:
        return ""
    prefers_excel_layout = should_prefer_excel_for_issuer(issuer)
    source_html = "<div style='margin-top:18px;'>"
    for block in blocks:
        source_html += "<div style='margin:10px 0;'>"
        is_excel_block = bool(block.get("excel_files")) and bool(block.get("html_fragment"))
        if block.get("pdf_images"):
            for image_path in block["pdf_images"]:
                try:
                    cid = attach_inline_image(mail, image_path)
                    source_html += (
                        f"<div style='margin:10px 0;'>"
                        f"<img src='cid:{cid}' style='max-width:100%; border:1px solid #CCCCCC;'>"
                        f"</div>"
                    )
                except Exception as e:
                    source_html += f"<p>[PDF 圖片插入失敗: {html.escape(str(e))}]</p>"
        else:
            if is_excel_block and not prefers_excel_layout:
                source_html += "<p>&nbsp;</p>"
            source_html += block.get("html_fragment", "")
        source_html += "</div>"
    source_html += "</div>"
    return source_html


def get_jpm_settle_date_from_outlook(start_date, end_date, isin):
    """JPMorgan 專用：從 Outlook 信件的 Excel 附件中找出指定 ISIN 的 Settle Date。"""
    isin_upper = str(isin or "").strip().upper()
    temp_dir = os.path.join(DRAFT_TEMP_PATH, f"jpm_settle_{uuid.uuid4().hex}")
    os.makedirs(temp_dir, exist_ok=True)
    try:
        outlook_ns = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
        target_folder = get_target_outlook_folder(
            outlook_namespace=outlook_ns,
            account_name=OUTLOOK_ACCOUNT,
            target_folder_name=TARGET_FOLDER_NAME
        )
        messages = get_outlook_messages(target_folder, start_date, end_date)
    except Exception as e:
        print(f"[JPM Settle Date] 無法存取 Outlook: {e}")
        shutil.rmtree(temp_dir, ignore_errors=True)
        return ""
    try:
        for msg in messages:
            try:
                if msg.Class != 43:
                    continue
                mail_date = msg.ReceivedTime.date()
                if mail_date < start_date or mail_date > end_date:
                    continue
                try:
                    sender_email = msg.SenderEmailAddress
                except Exception:
                    sender_email = ""
                sender_name = getattr(msg, "SenderName", "")
                extracted_company = extract_company_domain(sender_email, sender_name)
                if not any(k in extracted_company.upper() for k in ["JPM", "JPMORGAN", "CHASE"]):
                    continue
                blob = "\n".join([
                    str(getattr(msg, "Subject", "")),
                    str(getattr(msg, "Body", ""))
                ]).upper()
                if isin_upper and isin_upper not in blob:
                    att_names = []
                    if msg.Attachments.Count > 0:
                        for i in range(1, msg.Attachments.Count + 1):
                            try:
                                att_names.append(str(msg.Attachments.Item(i).FileName))
                            except Exception:
                                pass
                    if not any(isin_upper in x.upper() for x in att_names):
                        continue
                if msg.Attachments.Count <= 0:
                    continue
                for i in range(1, msg.Attachments.Count + 1):
                    try:
                        att = msg.Attachments.Item(i)
                        filename = str(att.FileName or "")
                        if not filename.lower().endswith((".xlsx", ".xlsm", ".xls")):
                            continue
                        excel_path = os.path.join(temp_dir, filename)
                        att.SaveAsFile(excel_path)
                        password = get_excel_password("JPM", filename)
                        temp_decrypted_path = None
                        try:
                            try:
                                all_sheets = pd.read_excel(excel_path, sheet_name=None, header=None, dtype=str, engine='openpyxl', keep_default_na=False)
                            except Exception:
                                if not password:
                                    continue
                                temp_decrypted_path = decrypt_excel_to_tempfile(excel_path, password)
                                all_sheets = pd.read_excel(temp_decrypted_path, sheet_name=None, header=None, dtype=str, engine='openpyxl', keep_default_na=False)
                            for sheet_name, df in all_sheets.items():
                                df = df.fillna("").astype(str)
                                header_idx = None
                                for idx, row in df.iterrows():
                                    row_vals_norm = [str(v).strip().upper() for v in row.values]
                                    has_isin = "ISIN" in row_vals_norm
                                    has_settle = any(k in row_vals_norm for k in ["SETTLE DATE", "SETTLEMENT DATE", "CASH SETTLEMENT DATE"])
                                    if has_isin and has_settle:
                                        header_idx = idx
                                        break
                                if header_idx is None:
                                    continue
                                header_vals = [str(v).strip() for v in df.iloc[header_idx].values]
                                isin_col_idx = None
                                settle_col_idx = None
                                for ci, h in enumerate(header_vals):
                                    h_upper = h.upper()
                                    if h_upper == "ISIN":
                                        isin_col_idx = ci
                                    if h_upper in ("SETTLE DATE", "SETTLEMENT DATE", "CASH SETTLEMENT DATE"):
                                        settle_col_idx = ci
                                if isin_col_idx is None or settle_col_idx is None:
                                    continue
                                for row_i in range(header_idx + 1, len(df)):
                                    row_vals = [str(v).strip() for v in df.iloc[row_i].values]
                                    if all(v in ("", "nan", "None") for v in row_vals):
                                        break
                                    cell_isin = row_vals[isin_col_idx] if isin_col_idx < len(row_vals) else ""
                                    if cell_isin.upper() == isin_upper:
                                        raw_date = row_vals[settle_col_idx] if settle_col_idx < len(row_vals) else ""
                                        if raw_date and raw_date not in ("", "nan", "None", "N/A"):
                                            try:
                                                parsed = pd.to_datetime(raw_date, dayfirst=False, errors='raise')
                                                return parsed.strftime("%Y-%m-%d")
                                            except Exception:
                                                return raw_date
                        finally:
                            if temp_decrypted_path and os.path.exists(temp_decrypted_path):
                                try:
                                    os.remove(temp_decrypted_path)
                                except Exception:
                                    pass
                    except Exception as e:
                        print(f"[JPM Settle Date] 附件處理失敗: {e}")
            except Exception as e:
                print(f"[JPM Settle Date] 信件處理失敗: {e}")
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)
    return ""


_eln_extracted_strike_cache = None


def normalize_equity_ticker_for_lookup(ticker):
    """將 Bloomberg ticker / RIC / 上手信文字統一成股票代碼主體，例如 MSFT.OQ、MSFT UW 都轉成 MSFT。"""
    text = str(ticker or "").strip().upper()
    if not text or text in ("NAN", "0", "相關股票"):
        return ""
    if ":" in text:
        text = text.split(":", 1)[1].strip()
    text = re.sub(r"\([^)]*\)", " ", text)
    text = re.sub(r"\b(EQUITY|REUTERS|RIC|US|UW|UN|UQ|UR|UT|UW)\b", " ", text)
    parts = re.findall(r"[A-Z0-9]+(?:[./][A-Z0-9]+)?", text)
    for part in parts:
        base = re.split(r"[./]", part)[0]
        if base and base not in {"USD", "HKD", "TWD", "JPY", "EUR", "GBP"}:
            return base
    return ""


def is_valid_strike_value(value):
    text = str(value or "").strip()
    if not text or text.lower() in ("nan", "0", "none"):
        return False
    try:
        return float(text.replace(",", "")) != 0
    except (ValueError, TypeError):
        return True


def _load_eln_extracted_data():
    """載入 ELN_Extracted_Data.xlsx，回傳 list of (isin, ticker_norm, strike)。僅首次呼叫時讀檔。"""
    global _eln_extracted_strike_cache
    if _eln_extracted_strike_cache is not None:
        return _eln_extracted_strike_cache
    _eln_extracted_strike_cache = []
    try:
        df_ext = pd.read_excel(ELN_EXTRACTED_DATA_PATH, header=0, dtype=str)
        for _, row in df_ext.iterrows():
            vals = row.values
            if len(vals) >= 7:
                row_isin   = str(vals[1]).strip()   # 第2欄 ISIN
                row_ticker = str(vals[3]).strip()   # 第4欄 (D欄) Ticker
                row_strike = str(vals[6]).strip()   # 第7欄 Strike_Price
                if row_isin and row_isin.lower() not in ('nan', ''):
                    ticker_norm = normalize_equity_ticker_for_lookup(row_ticker)
                    _eln_extracted_strike_cache.append((row_isin.upper(), ticker_norm, row_strike))
        print(f"[INFO] ELN_Extracted_Data.xlsx 載入完成，共 {len(_eln_extracted_strike_cache)} 筆資料")
    except Exception as e:
        print(f"[警告] 無法載入 ELN_Extracted_Data.xlsx: {e}")
    return _eln_extracted_strike_cache


def lookup_strike_price(isin, ticker=""):
    """從 ELN_Extracted_Data.xlsx 查詢履約價：優先 ISIN+Ticker 雙重比對，其次 ISIN 單獨比對。"""
    rows = _load_eln_extracted_data()
    isin_upper  = str(isin   or "").strip().upper()
    ticker_norm = normalize_equity_ticker_for_lookup(ticker)
    # 優先：ISIN + Ticker 都符合
    if ticker_norm:
        for row_isin, row_ticker, row_strike in rows:
            if row_isin == isin_upper and row_ticker == ticker_norm:
                if is_valid_strike_value(row_strike):
                    return row_strike
    # 次要：ISIN 單獨符合，但多標的 / 多履約價時不任意拿第一筆，避免接錯標的履約價
    isin_matches = []
    for row_isin, row_ticker, row_strike in rows:
        if row_isin == isin_upper:
            if is_valid_strike_value(row_strike):
                isin_matches.append(row_strike)
    distinct_strikes = []
    for strike in isin_matches:
        key = str(strike).strip().replace(",", "")
        if key not in distinct_strikes:
            distinct_strikes.append(key)
    if len(distinct_strikes) == 1:
        return isin_matches[0]
    return ""


def normalize_subject_text(value):
    if pd.isna(value):
        return ""
    text = str(value).replace("\r", " ").replace("\n", " ").strip()
    text = re.sub(r"\s+", " ", text)
    if text.lower() in {"", "nan", "none"}:
        return ""
    return text


def pick_best_product_name(metadata, group_df):
    product_col = "\u5546\u54c1\u540d\u7a31"
    candidates = []

    if product_col in group_df.columns:
        for value in group_df[product_col].tolist():
            normalized = normalize_subject_text(value)
            if normalized and normalized not in candidates:
                candidates.append(normalized)

    fallback = normalize_subject_text(metadata.get(product_col, ""))
    if fallback and fallback not in candidates:
        candidates.append(fallback)

    if not candidates:
        return "BEN / FCN"

    return max(candidates, key=lambda text: (len(text), text.count("（") + text.count("("), text))


def unique_preserve_order(values):
    result = []
    for value in values:
        cleaned = normalize_subject_text(value)
        if cleaned and cleaned not in result:
            result.append(cleaned)
    return result


def normalize_branch_name(value):
    text = normalize_subject_text(value)
    if not text:
        return ""
    if "\u53f0\u4e2d" in text:
        return "\u53f0\u4e2d"
    if "\u9ad8\u96c4" in text:
        return "\u9ad8\u96c4"
    if "\u7e3d\u516c\u53f8" in text:
        return "\u7e3d\u516c\u53f8"
    return text


def get_outlook_recipients_by_branch(group_df):
    branch_col = "\u5206\u516c\u53f8"
    branch_order = {
        "\u7e3d\u516c\u53f8": 0,
        "\u53f0\u4e2d": 1,
        "\u9ad8\u96c4": 2,
    }
    routing_map = {
        ("\u7e3d\u516c\u53f8",): {
            "to": [
                "\u696d\u52d9\u52a9\u7406(SA)",
                "\u5546\u54c1\u4f01\u5283\u5ba4(Product Planning office)",
            ],
            "cc": [
                "\u7d50\u7b97\u4f5c\u696d\u90e8(Operations)",
            ],
        },
        ("\u53f0\u4e2d",): {
            "to": [
                "Ivy Lin \u6797\u5b5f\u7487",
                "\u7d50\u7b97\u4f5c\u696d\u90e8\u53f0\u4e2d(Operations)",
                "\u5546\u54c1\u4f01\u5283\u5ba4(Product Planning office)",
            ],
            "cc": [
                "\u7d50\u7b97\u4f5c\u696d\u90e8(Operations)",
                "\u696d\u52d9\u52a9\u7406(SA)",
            ],
        },
        ("\u9ad8\u96c4",): {
            "to": [
                "Vicky Wu \u5433\u6620\u67d4",
                "\u7d50\u7b97\u4f5c\u696d\u90e8\u9ad8\u96c4(Operations)",
                "\u5546\u54c1\u4f01\u5283\u5ba4(Product Planning office)",
            ],
            "cc": [
                "\u7d50\u7b97\u4f5c\u696d\u90e8(Operations)",
                "\u696d\u52d9\u52a9\u7406(SA)",
            ],
        },
        ("\u7e3d\u516c\u53f8", "\u53f0\u4e2d"): {
            "to": [
                "Ivy Lin \u6797\u5b5f\u7487",
                "\u7d50\u7b97\u4f5c\u696d\u90e8\u53f0\u4e2d(Operations)",
                "\u696d\u52d9\u52a9\u7406(SA)",
                "\u5546\u54c1\u4f01\u5283\u5ba4(Product Planning office)",
            ],
            "cc": [
                "\u7d50\u7b97\u4f5c\u696d\u90e8(Operations)",
                "\u696d\u52d9\u52a9\u7406(SA)",
            ],
        },
        ("\u7e3d\u516c\u53f8", "\u9ad8\u96c4"): {
            "to": [
                "Vicky Wu \u5433\u6620\u67d4",
                "\u7d50\u7b97\u4f5c\u696d\u90e8\u9ad8\u96c4(Operations)",
                "\u696d\u52d9\u52a9\u7406(SA)",
                "\u5546\u54c1\u4f01\u5283\u5ba4(Product Planning office)",
            ],
            "cc": [
                "\u7d50\u7b97\u4f5c\u696d\u90e8(Operations)",
                "\u696d\u52d9\u52a9\u7406(SA)",
            ],
        },
        ("\u53f0\u4e2d", "\u9ad8\u96c4"): {
            "to": [
                "Ivy Lin \u6797\u5b5f\u7487",
                "\u7d50\u7b97\u4f5c\u696d\u90e8\u53f0\u4e2d(Operations)",
                "Vicky Wu \u5433\u6620\u67d4",
                "\u7d50\u7b97\u4f5c\u696d\u90e8\u9ad8\u96c4(Operations)",
                "\u696d\u52d9\u52a9\u7406(SA)",
                "\u5546\u54c1\u4f01\u5283\u5ba4(Product Planning office)",
            ],
            "cc": [
                "\u7d50\u7b97\u4f5c\u696d\u90e8(Operations)",
                "\u696d\u52d9\u52a9\u7406(SA)",
            ],
        },
    }

    branches = []
    if branch_col in group_df.columns:
        for value in group_df[branch_col].tolist():
            branch_name = normalize_branch_name(value)
            if branch_name and branch_name not in branches:
                branches.append(branch_name)

    branches = sorted(branches, key=lambda item: (branch_order.get(item, 99), item))
    route_key = tuple(branches)

    if route_key in routing_map:
        routing = routing_map[route_key]
        return route_key, routing["to"], routing["cc"]

    has_head_office = "\u7e3d\u516c\u53f8" in branches
    has_taichung = "\u53f0\u4e2d" in branches
    has_kaohsiung = "\u9ad8\u96c4" in branches

    if not branches:
        default_routing = routing_map[("\u7e3d\u516c\u53f8",)]
        return ("\u7e3d\u516c\u53f8",), default_routing["to"], default_routing["cc"]
    if not (has_head_office or has_taichung or has_kaohsiung):
        default_routing = routing_map[("\u7e3d\u516c\u53f8",)]
        return route_key, default_routing["to"], default_routing["cc"]

    to_list = []
    cc_list = ["\u7d50\u7b97\u4f5c\u696d\u90e8(Operations)"]

    if has_taichung:
        to_list.extend([
            "Ivy Lin \u6797\u5b5f\u7487",
            "\u7d50\u7b97\u4f5c\u696d\u90e8\u53f0\u4e2d(Operations)",
        ])
    if has_kaohsiung:
        to_list.extend([
            "Vicky Wu \u5433\u6620\u67d4",
            "\u7d50\u7b97\u4f5c\u696d\u90e8\u9ad8\u96c4(Operations)",
        ])

    if has_head_office and not (has_taichung or has_kaohsiung):
        to_list.append("\u696d\u52d9\u52a9\u7406(SA)")
    elif has_head_office or sum([has_taichung, has_kaohsiung]) >= 2:
        to_list.append("\u696d\u52d9\u52a9\u7406(SA)")
        cc_list.append("\u696d\u52d9\u52a9\u7406(SA)")
    elif has_taichung or has_kaohsiung:
        cc_list.append("\u696d\u52d9\u52a9\u7406(SA)")

    if has_head_office or has_taichung or has_kaohsiung:
        to_list.append("\u5546\u54c1\u4f01\u5283\u5ba4(Product Planning office)")

    return route_key, unique_preserve_order(to_list), unique_preserve_order(cc_list)


def generate_outlook_draft(isin, category, metadata, group_df, start_date, end_date):
    """根據判斷結果與客戶明細表，產出 Outlook 草稿信件並顯示。"""
    try:
        outlook = win32com.client.Dispatch("Outlook.Application")
        mail = outlook.CreateItem(0)
        branch_route, to_recipients, cc_recipients = get_outlook_recipients_by_branch(group_df)
        branch_route_label = " + ".join(branch_route) if branch_route else "\u7e3d\u516c\u53f8"
        if to_recipients:
            mail.To = "; ".join(to_recipients)
        if cc_recipients:
            mail.CC = "; ".join(cc_recipients)
        used_strike = ""  # 本次查詢到的履約價，供呼叫端寫回 Excel
        obs_date = str(metadata.get("ObsDate", metadata.get("Date", ""))).split(' ')[0].strip()
        pay_date = str(metadata.get("Value Day/Payment Date", metadata.get("Payment Date", ""))).split(' ')[0].strip()
        issuer = str(metadata.get("Company", "發行機構")).upper()
        issuer_norm = normalize_company_key(issuer)
        if issuer_norm in {"JPM"} and (not pay_date or pay_date.upper() in ("CHECK", "CHECK FILE", "N/A", "NAN", "")):
            settle = get_jpm_settle_date_from_outlook(start_date, end_date, isin)
            if settle:
                pay_date = settle
        raw_value = str(metadata.get("Value/Raw Data", ""))
        raw_trade_date = str(metadata.get('下單日期', 'YYYY/MM/DD'))
        trade_date = raw_trade_date.split(' ')[0].strip()
        product_name = pick_best_product_name(metadata, group_df)
        desired_cols = ['投資人帳號', '分公司', 'ISIN', 'Company', '本日贖回單位數', '幣別']
        actual_cols = [col for col in desired_cols if col in group_df.columns]
        if not actual_cols:
            candidate_cols = [c for c in list(group_df.columns) if group_df[c].notna().any()]
            actual_cols = candidate_cols[:6]
        table_df = group_df[actual_cols].copy() if actual_cols else pd.DataFrame()
        position_table_html = create_html_table(table_df)
        style = "font-family: 'Microsoft JhengHei', sans-serif; color: #5C2483; font-size: 11pt;"
        html_body = f"<div style=\"{style}\"><p>親愛的遠智證券客戶，您好。</p>"
        template_name = ""
        if category == 'conversion':
            template_name = "轉換通知"
            stock_ticker = raw_value.replace("換股標的:", "").strip() if "換股標的" in raw_value else "相關股票"
            if issuer_norm == "BARCLAYS" and (stock_ticker == "?賊??∠巨" or looks_like_isin_only(stock_ticker)):
                recovered_ticker = lookup_barclays_ticker_from_outlook(start_date, end_date, isin)
                if recovered_ticker:
                    stock_ticker = recovered_ticker
            if issuer_norm == "BARCLAYS" and (not normalize_equity_ticker_for_lookup(stock_ticker) or looks_like_isin_only(stock_ticker)):
                recovered_ticker = lookup_barclays_ticker_from_outlook(start_date, end_date, isin)
                if recovered_ticker:
                    stock_ticker = recovered_ticker
            denom_str = str(metadata.get('Denom', '10,000')).strip()
            # Nomura: Denom 欄位存的是總名目本金，每單位固定為 10,000
            if "NOMURA" in issuer.upper():
                denom_str = '10,000'
            strike_str = str(metadata.get('Strike', '0')).strip()
            # 從 ELN_Extracted_Data.xlsx 以 ISIN + Ticker 查詢履約價 (轉換股票專用)。
            # UBS RIC 接股通知本身會提供轉換價，優先保留上手信 Strike，避免 ISIN fallback 接到舊條款價。
            extracted_strike = ""
            if not (issuer_norm == "UBS" and is_valid_strike_value(strike_str)):
                extracted_strike = lookup_strike_price(isin, stock_ticker)
            if is_valid_strike_value(extracted_strike):
                strike_str = extracted_strike
                used_strike = extracted_strike  # 記錄查詢到的履約價
            shares_str = str(metadata.get('Shares', '0')).strip()
            cash_str = str(metadata.get('FractionalCash', '0')).strip()
            currency_str = str(metadata.get('幣別', 'USD')).strip()
            if "MORGAN" in issuer and "STANLEY" in issuer:
                total_units = pd.to_numeric(group_df.get('本日贖回單位數', pd.Series([1])), errors='coerce').sum()
                if pd.isna(total_units) or total_units <= 0:
                    total_units = 1
                try:
                    d_val = float(denom_str.replace(",", ""))
                    if d_val > 0:
                        total_units = total_units / (d_val / 10000.0)
                except Exception:
                    pass
                try:
                    s_val = float(shares_str.replace(",", ""))
                    calculated_share = s_val / float(total_units)
                    shares_str = f"{calculated_share:g}"
                except Exception:
                    pass
                cash_match = re.search(r"^([\d\.,]+)\s*\(\+Coupon:([\d\.,]+)\)", cash_str)
                if cash_match:
                    try:
                        t_payable = float(cash_match.group(1).replace(",", ""))
                        coupon = float(cash_match.group(2).replace(",", ""))
                        per_unit_cash = (t_payable / float(total_units)) - coupon
                        cash_str = f"{per_unit_cash:g}"
                    except Exception:
                        pass
                else:
                    try:
                        t_payable = float(cash_str.replace(",", ""))
                        per_unit_cash = t_payable / float(total_units)
                        cash_str = f"{per_unit_cash:g}"
                    except Exception:
                        pass
            try:
                d_val = float(denom_str.replace(",", ""))
                denom_fmt = f"{int(d_val):,}" if d_val.is_integer() else f"{d_val:,}"
            except Exception:
                denom_fmt = denom_str
            # 分割/反分割偵測：Denom/Strike 的整數部分應等於接整股數
            if (strike_str and strike_str not in ("0", "nan")
                    and shares_str and shares_str not in ("0", "nan")):
                try:
                    expected_shares = int(float(denom_str.replace(",", "")) / float(strike_str.replace(",", "")))
                    actual_shares = int(float(shares_str.replace(",", "")))
                    if expected_shares != actual_shares:
                        print(f"\n{'!' * 60}")
                        print(f"[警示] 偵測到可能的股票分割/反分割活動！")
                        print(f"  ISIN          : {isin}")
                        print(f"  Denom         : {denom_str}")
                        print(f"  履約價 Strike : {strike_str}")
                        print(f"  Denom/Strike  : {expected_shares} (整數部分)")
                        print(f"  實際接整股數  : {actual_shares}")
                        print(f"  → 兩者不符，請確認是否發生分割 (Split) 或反分割 (Reverse Split) 活動！")
                        print(f"{'!' * 60}\n")
                except (ValueError, TypeError):
                    pass
            if strike_str and strike_str != "0" and strike_str != "nan":
                detail_msg = f"每單位{denom_fmt} => {denom_fmt}/{strike_str}, "
            else:
                detail_msg = f"每單位{denom_fmt} => {denom_fmt}/[未取得履約價], "
            detail_msg += f"每單位收{shares_str}整股{stock_ticker}"
            if cash_str and cash_str != "0" and cash_str != "nan":
                detail_msg += f"每單位零股等值約{cash_str}{currency_str}"
            mail.Subject = f"【結構型商品】轉換通知 {trade_date.replace('-', '/')} {issuer} {product_name} {isin}"
            html_body += f"""
            <p>您於{trade_date.replace('-', '/')}所承做的{product_name} (ISIN Code : {html.escape(str(isin))})，因{html.escape(obs_date)}連結標的價格低於接股價格，
            發行機構將於{html.escape(pay_date)}撥出{html.escape(stock_ticker)}股票，但實際交割日期、金額或/與股數以{html.escape(issuer)}發出最後交割指示為準，特此通知。</p>
            <p style="color: #FF0000; font-weight: bold;">{html.escape(detail_msg)}。</p>
            """.strip()
        elif category == 'early':
            template_name = "提前買回通知"
            mail.Subject = f"【結構型商品】提前買回通知 {trade_date.replace('-', '/')} {issuer} {product_name} {isin}"
            html_body += f"""
            <p>您於{trade_date.replace('-', '/')}承做的{product_name} (ISIN Code : {html.escape(str(isin))})，因{html.escape(obs_date)}的連結標的觸發自動提前出場事件，
            交割日為 {html.escape(pay_date)}，但實際交割日期、金額或/與股數以{html.escape(issuer)}發出最後交割指示為準，特此通知。</p>
            <p>若您尚有相關疑問，歡迎致電您的服務人員垂詢，謝謝。</p>
            """.strip()
        elif category == 'maturity':
            template_name = "到期通知"
            mail.Subject = f"【結構型商品】到期通知 {trade_date.replace('-', '/')} {issuer} {product_name} {isin}"
            html_body += f"""
            <p>您於{trade_date.replace('-', '/')}承做的{product_name} (ISIN Code : {html.escape(str(isin))})，因{html.escape(obs_date)}最後評價日且標的價格大於下限價，
            交割日為 {html.escape(pay_date)}，但實際交割日期、金額或/與股數以{html.escape(issuer)}發出最後交割指示為準，特此通知。</p>
            <p>若您尚有相關疑問，歡迎致電您的服務人員垂詢，謝謝。</p>
            """.strip()
        else:
            print(f"略過 ISIN {isin}: 不支援的 category = {category}")
            return False, ""
        event_type_text = str(metadata.get('Event Type', '')).strip()
        source_content_html = build_source_content_html(
            mail=mail,
            start_date=start_date,
            end_date=end_date,
            issuer=issuer,
            isin=isin,
            target_category=category,
            target_event_type=event_type_text
        )
        html_body += position_table_html
        if source_content_html:
            html_body += "<p>&nbsp;</p>"
            html_body += source_content_html
        html_body += "</div>"
        mail.HTMLBody = html_body
        mail.Save()
        mail.Display()
        print(f"成功建立草稿 [{template_name}] | ISIN: {isin} | 分公司規則: {branch_route_label} (包含 {len(group_df)} 筆明細)")
        return True, used_strike
    except Exception as e:
        print(f"建立草稿失敗 (ISIN {isin}): {e}")
        return False, ""


def step4_generate_outlook_drafts(start_date, end_date, final_file):
    """
    步驟 4 (ELN4)：讀取 ELN_Final_Output → 依 ISIN 分組判斷事件類型 → 產生 Outlook 草稿
    """
    print("\n=== 步驟 4：產生 Outlook 草稿開始 ===")

    if not os.path.exists(final_file):
        print(f"錯誤：找不到步驟2的輸出檔 '{final_file}'")
        return

    try:
        df = pd.read_excel(final_file)
        df.columns = df.columns.str.strip()
    except Exception as e:
        print(f"Excel 讀取失敗: {e}")
        return

    if 'Event Type' not in df.columns or 'ISIN' not in df.columns:
        print("錯誤：Excel 中找不到 'Event Type' 或 'ISIN' 欄位")
        return

    has_raw_data_col = 'Value/Raw Data' in df.columns
    process_df = build_position_process_df(df)
    if process_df.empty and '成交種類' not in df.columns:
        process_df = df.copy()

    grouped = process_df.groupby('ISIN')
    stats = {'early': 0, 'maturity': 0, 'conversion': 0, 'skip': 0}
    isin_strike_map = {}  # 記錄本次查詢到的 ISIN → Strike，供寫回 Excel 用

    # 確保日期是 date 物件（pipeline 傳入的是 date，若直接呼叫傳入字串也能相容）
    start_date_obj = start_date if isinstance(start_date, datetime.date) else datetime.datetime.strptime(str(start_date), "%Y-%m-%d").date()
    end_date_obj = end_date if isinstance(end_date, datetime.date) else datetime.datetime.strptime(str(end_date), "%Y-%m-%d").date()

    print(f"\n共發現 {len(grouped)} 檔獨立 ISIN 商品，開始生成 Outlook 草稿...\n")

    for isin, group_df in grouped:
        first_row = group_df.iloc[0]
        category = classify_row(first_row, has_raw_data_col)
        if category is None:
            print(f"跳過 ISIN {isin}: 無法判定 Event Type ({first_row.get('Event Type', '')})")
            stats['skip'] += 1
            continue

        success, used_strike = generate_outlook_draft(
            isin=isin,
            category=category,
            metadata=first_row,
            group_df=group_df,
            start_date=start_date_obj,
            end_date=end_date_obj
        )
        if success:
            stats[category] += 1
            if category == 'conversion' and used_strike and used_strike not in ('0', 'nan', ''):
                isin_strike_map[str(isin).strip()] = used_strike

    print("\n" + "=" * 50)
    print("步驟4執行完成！統計結果：")
    print(f" - 提前買回通知: {stats['early']} 封")
    print(f" - 到期通知:     {stats['maturity']} 封")
    print(f" - 轉換通知:     {stats['conversion']} 封")
    print(f" - 未分類/跳過: {stats['skip']} 封")
    print("=" * 50)

    # --- 將查詢到的 Strike 寫回 ELN_Final_Output 及 ELN_Report ---
    if isin_strike_map:
        def _write_strike_back(file_path, strike_map):
            if not os.path.exists(file_path):
                return
            try:
                df_w = pd.read_excel(file_path)
                df_w.columns = df_w.columns.str.strip()
                if 'ISIN' not in df_w.columns or 'Strike' not in df_w.columns:
                    return
                df_w['ISIN'] = df_w['ISIN'].astype(str).str.strip()
                df_w['Strike'] = pd.to_numeric(df_w['Strike'], errors='coerce')  # int64 → float64
                updated = 0
                for isin_key, strike_val in strike_map.items():
                    mask = df_w['ISIN'] == isin_key
                    if mask.any():
                        try:
                            df_w.loc[mask, 'Strike'] = float(strike_val)
                        except (ValueError, TypeError):
                            df_w['Strike'] = df_w['Strike'].astype(object)
                            df_w.loc[mask, 'Strike'] = strike_val
                        updated += mask.sum()
                df_w.to_excel(file_path, index=False)
                print(f"[Strike回寫] {os.path.basename(file_path)} 更新 {updated} 筆")
            except Exception as e:
                print(f"[Strike回寫失敗] {file_path}: {e}")

        _write_strike_back(final_file, isin_strike_map)
        report_file = build_report_output_path(start_date_obj, end_date_obj)
        _write_strike_back(report_file, isin_strike_map)


# ==========================================
# 主流程 (Pipeline)
# ==========================================

def main():
    """一次執行三個步驟：Outlook擷取 → 合併統計表 → 產生Word通知書"""
    print("=" * 60)
    print("      ELN 全流程自動化 Pipeline")
    print("=" * 60)

    # 只需輸入一次日期，三個步驟共用
    start_date, end_date = get_user_date_range()

    # 步驟 1：從 Outlook 擷取信件並輸出 ELN_Report
    report_file = step1_extract_from_outlook(start_date, end_date)
    if report_file is None:
        print("\n步驟1失敗，Pipeline 中止。")
        return

    # 步驟 2：合併 ELN統計表 並輸出 ELN_Final_Output
    final_file = step2_merge_with_stats(start_date, end_date, report_file)
    if final_file is None:
        print("\n步驟2失敗，Pipeline 中止。")
        return

    # 步驟 3：產生 Word 通知書
    step3_generate_documents(start_date, end_date, final_file)

    # 步驟 4：產生 Outlook 草稿
    step4_generate_outlook_drafts(start_date, end_date, final_file)

    print("\n=== Pipeline 全部完成 ===")


if __name__ == "__main__":
    main()